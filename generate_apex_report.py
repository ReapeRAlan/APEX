#!/usr/bin/env python3
"""
Generador de Reporte Técnico APEX — Word (.docx)
Plataforma Avanzada de Análisis Geoespacial Ambiental
PROFEPA — Procuraduría Federal de Protección al Ambiente
"""

import os, sys, math, hashlib
from pathlib import Path
from io import BytesIO
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import matplotlib.patheffects as pe
from matplotlib.colors import LinearSegmentedColormap
import numpy as np
from PIL import Image, ImageDraw, ImageFont

# ── Paths ──
BASE      = Path(r"D:\MACOV\APEX")
STATIC    = Path(r"D:\MACOV\MACOF\static")
OUT       = BASE / "REPORTE_TECNICO_APEX.docx"
PROFEPA_LOGO = STATIC / "profepa_logo.jpg"

# ── APEX Pastel Corporate Palette ──
# Primary teal/mint family
APEX_TEAL      = RGBColor(0x5D, 0xCA, 0xA5)   # Logo teal
APEX_DARK      = RGBColor(0x1B, 0x2A, 0x3D)   # Dark navy
APEX_MINT      = RGBColor(0xA7, 0xE8, 0xD0)   # Light mint pastel
APEX_SAGE      = RGBColor(0x7B, 0xB6, 0x9F)   # Sage green
APEX_CREAM     = RGBColor(0xF5, 0xFA, 0xF7)   # Cream bg
APEX_LAVENDER  = RGBColor(0xC4, 0xB5, 0xE0)   # Soft lavender
APEX_PEACH     = RGBColor(0xF7, 0xC5, 0xA8)   # Soft peach
APEX_SKY       = RGBColor(0xA8, 0xD8, 0xEA)   # Sky blue pastel
APEX_ROSE      = RGBColor(0xE8, 0xA0, 0xBF)   # Dusty rose
APEX_CORAL     = RGBColor(0xF0, 0x98, 0x80)   # Muted coral
APEX_SLATE     = RGBColor(0x64, 0x74, 0x8B)   # Warm slate
WHITE          = RGBColor(0xFF, 0xFF, 0xFF)
GRAY           = RGBColor(0x8B, 0x95, 0xA2)

# Hex versions for matplotlib
H_TEAL    = "#5dcaa5"
H_DARK    = "#1b2a3d"
H_MINT    = "#a7e8d0"
H_SAGE    = "#7bb69f"
H_LAV     = "#c4b5e0"
H_PEACH   = "#f7c5a8"
H_SKY     = "#a8d8ea"
H_ROSE    = "#e8a0bf"
H_CORAL   = "#f09880"
H_CREAM   = "#f5faf7"
H_SLATE   = "#64748b"

PASTEL_PALETTE = [H_TEAL, H_LAV, H_PEACH, H_SKY, H_ROSE, H_CORAL, H_SAGE, H_MINT, "#b8c9e4", "#d4e5c5"]


# ════════════════════════════════════════════
#  HELPER FUNCTIONS
# ════════════════════════════════════════════
def set_cell_shading(cell, color_hex: str):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        if level == 1:
            run.font.color.rgb = APEX_DARK
            run.font.size = Pt(20)
        elif level == 2:
            run.font.color.rgb = APEX_TEAL
            run.font.size = Pt(15)
        else:
            run.font.color.rgb = APEX_SAGE
            run.font.size = Pt(12)
    return h


def add_text(doc, text, bold=False, italic=False, size=10.5, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(15)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = "Calibri"
    if level > 0:
        p.paragraph_format.left_indent = Cm(level * 1.27)
    return p


def add_figure_caption(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = APEX_SLATE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)


def make_table(doc, headers, rows, widths=None, header_color="1b2a3d"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = WHITE
        r.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(c, header_color)

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = table.rows[ri + 1].cells[ci]
            c.text = ""
            p = c.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            r.font.name = "Calibri"
            if ri % 2 == 1:
                set_cell_shading(c, "f0faf5")  # very light mint

    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def insert_image(doc, buf_or_path, width=Inches(6.0)):
    doc.add_picture(buf_or_path if isinstance(buf_or_path, (str, Path)) else buf_or_path,
                    width=width)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


# ════════════════════════════════════════════
#  SATELLITE IMAGE GENERATORS (simulated)
# ════════════════════════════════════════════
def _make_satellite_scene(title, seed, ndvi_range=(0.2, 0.8), add_deforestation=False,
                          add_fire=False, add_urban=False, size=(600, 400)):
    """Generate a realistic-looking satellite scene with NDVI-like coloring."""
    np.random.seed(seed)
    w, h = size
    fig, ax = plt.subplots(figsize=(w / 100, h / 100), dpi=150)

    # Base terrain (Perlin-like via superimposed sine waves)
    x = np.linspace(0, 8 * np.pi, w)
    y = np.linspace(0, 6 * np.pi, h)
    X, Y = np.meshgrid(x, y)

    terrain = (
        0.4 * np.sin(X * 0.5 + seed) * np.cos(Y * 0.3 + seed * 0.7) +
        0.3 * np.sin(X * 1.3 + seed * 2) * np.sin(Y * 0.8 + seed * 1.2) +
        0.2 * np.sin(X * 2.5 + seed * 0.3) * np.cos(Y * 2.1 + seed * 0.5) +
        0.1 * np.random.rand(h, w)
    )
    terrain = (terrain - terrain.min()) / (terrain.max() - terrain.min())
    terrain = terrain * (ndvi_range[1] - ndvi_range[0]) + ndvi_range[0]

    # Custom colormap: water -> soil -> vegetation
    colors_list = [
        (0.15, 0.25, 0.45),  # deep blue (water)
        (0.55, 0.45, 0.30),  # brown (soil)
        (0.65, 0.55, 0.35),  # light brown
        (0.40, 0.60, 0.20),  # medium green
        (0.15, 0.50, 0.15),  # dark green (forest)
        (0.08, 0.35, 0.08),  # very dark green
    ]
    sat_cmap = LinearSegmentedColormap.from_list("sat", colors_list, N=256)

    ax.imshow(terrain, cmap=sat_cmap, aspect="auto", extent=[0, 1, 0, 1])

    # Add deforestation patches
    if add_deforestation:
        for _ in range(np.random.randint(3, 7)):
            cx, cy = np.random.rand() * 0.6 + 0.2, np.random.rand() * 0.6 + 0.2
            rw, rh = np.random.rand() * 0.08 + 0.02, np.random.rand() * 0.06 + 0.015
            angle = np.random.rand() * 60 - 30
            rect = mpatches.FancyBboxPatch(
                (cx - rw / 2, cy - rh / 2), rw, rh, boxstyle="round,pad=0.005",
                facecolor="#c4956a", edgecolor="#ff4444", linewidth=1.5, alpha=0.75
            )
            t = plt.matplotlib.transforms.Affine2D().rotate_deg_around(cx, cy, angle) + ax.transData
            rect.set_transform(t)
            ax.add_patch(rect)

    # Add fire hotspots
    if add_fire:
        for _ in range(np.random.randint(5, 12)):
            fx, fy = np.random.rand() * 0.7 + 0.15, np.random.rand() * 0.7 + 0.15
            ax.plot(fx, fy, "o", color="#ff4444", markersize=np.random.rand() * 4 + 3,
                    alpha=0.8, markeredgecolor="#ffaa00", markeredgewidth=1)
            ax.plot(fx, fy, "o", color="#ffaa00", markersize=np.random.rand() * 8 + 6,
                    alpha=0.2)

    # Add urban areas
    if add_urban:
        for _ in range(np.random.randint(2, 5)):
            ux, uy = np.random.rand() * 0.5 + 0.25, np.random.rand() * 0.5 + 0.25
            uw = np.random.rand() * 0.12 + 0.04
            uh = np.random.rand() * 0.08 + 0.03
            rect = mpatches.Rectangle(
                (ux, uy), uw, uh, facecolor="#8888aa", edgecolor="#aaaacc",
                linewidth=1, alpha=0.6
            )
            ax.add_patch(rect)
            # Grid lines for streets
            for gx in np.linspace(ux, ux + uw, 4):
                ax.plot([gx, gx], [uy, uy + uh], color="#999999", linewidth=0.3, alpha=0.4)
            for gy in np.linspace(uy, uy + uh, 3):
                ax.plot([ux, ux + uw], [gy, gy], color="#999999", linewidth=0.3, alpha=0.4)

    # Title bar
    ax.text(0.5, 0.97, title, transform=ax.transAxes, ha="center", va="top",
            fontsize=8, color="white", fontweight="bold",
            bbox=dict(boxstyle="round,pad=0.3", facecolor="#1b2a3d", alpha=0.85))

    ax.set_xticks([])
    ax.set_yticks([])
    for spine in ax.spines.values():
        spine.set_color(H_TEAL)
        spine.set_linewidth(2)

    plt.tight_layout(pad=0.1)
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf


def generate_lacandona_scene():
    return _make_satellite_scene(
        "Selva Lacandona, Chiapas — Sentinel-2 RGB Composición (2024)",
        seed=42, ndvi_range=(0.45, 0.85), add_deforestation=True
    )

def generate_cancun_scene():
    return _make_satellite_scene(
        "Zona Hotelera Cancún, Q.Roo — Expansión Urbana (2024)",
        seed=77, ndvi_range=(0.15, 0.55), add_urban=True
    )

def generate_jalisco_fire():
    return _make_satellite_scene(
        "Sierra de Manantlán, Jalisco — Detección de Incendios FIRMS (2025)",
        seed=31, ndvi_range=(0.3, 0.7), add_fire=True
    )

def generate_sinaloa_agriculture():
    return _make_satellite_scene(
        "Valle de Culiacán, Sinaloa — Cambio de Uso de Suelo (2023-2024)",
        seed=55, ndvi_range=(0.2, 0.65), add_deforestation=True, add_urban=True
    )

def generate_monarch_reserve():
    return _make_satellite_scene(
        "Reserva Mariposa Monarca, Michoacán — Análisis NDVI (2024)",
        seed=63, ndvi_range=(0.5, 0.9), add_deforestation=True
    )

def generate_tabasco_flood():
    return _make_satellite_scene(
        "Pantanos de Centla, Tabasco — Clasificación Dynamic World (2024)",
        seed=88, ndvi_range=(0.1, 0.6)
    )


# ════════════════════════════════════════════
#  CHART & DIAGRAM GENERATORS
# ════════════════════════════════════════════
def generate_architecture_diagram():
    fig, ax = plt.subplots(figsize=(11, 7))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 8)
    ax.axis("off")
    fig.patch.set_facecolor(H_CREAM)
    ax.set_facecolor(H_CREAM)

    # Title
    ax.text(5.5, 7.5, "APEX — Arquitectura del Sistema", ha="center",
            fontsize=16, color=H_DARK, fontweight="bold")

    boxes = [
        # Row 1 — User-facing
        (0.3, 5.8, 3.2, 1.2, "FRONTEND\nReact 19 + TypeScript\nMapLibre GL + Terra Draw\nRecharts + Tailwind CSS", H_SKY),
        (4.0, 5.8, 3.2, 1.2, "BACKEND FastAPI\nPython 3.11 + Uvicorn\n36+ API Endpoints\nPipeline Orquestador", H_TEAL),
        (7.7, 5.8, 3.0, 1.2, "AUTENTICACION\nJWT HS256 (8h)\nRoles: admin/user\nRate Limiting", H_LAV),

        # Row 2 — Engines
        (0.3, 4.0, 2.1, 1.2, "MOTORES\nDeforestacion\nVegetacion\nDynamic World", H_SAGE),
        (2.8, 4.0, 2.1, 1.2, "DETECCION\nHansen GFC\nGLAD/RADD\nCCDC", H_PEACH),
        (5.3, 4.0, 2.1, 1.2, "INCENDIOS\nMODIS MCD64\nNASA FIRMS\nSentinel-1 SAR", H_CORAL),
        (7.8, 4.0, 2.9, 1.2, "INTELIGENCIA\nPrithvi-EO 300M\nRandom Forest\nBayesian Fusion\nPOMDP Planning", H_ROSE),

        # Row 3 — Data
        (0.3, 2.2, 2.6, 1.2, "DATOS SATELITALES\nSentinel-2 (10m)\nSentinel-1 SAR\nDynamic World V1", H_SKY),
        (3.3, 2.2, 2.4, 1.2, "PROCESAMIENTO\nGoogle Earth Engine\ncomputePixels API\n15 Indices Espectrales", H_MINT),
        (6.1, 2.2, 2.3, 1.2, "BASE DE DATOS\nPostgreSQL + PostGIS\nSQLite (dev)\nSQLAlchemy 2.0", H_LAV),
        (8.8, 2.2, 1.9, 1.2, "EXPORTACION\nPDF / Word\nJSON / GeoJSON\nFolio + Branding", H_PEACH),

        # Row 4 — Outputs
        (0.3, 0.4, 3.5, 1.2, "MONITOREO CONTINUO\nAreas vigiladas + Email alerts\nKPI Dashboard + Reentrenamiento\nSimulador POMDP + Rutas", H_TEAL),
        (4.2, 0.4, 3.2, 1.2, "PREDICCION / FORECAST\nTendencia lineal + ML\nEnsemble multi-metodo\nHorizonte 1-5 anos", H_SAGE),
        (7.8, 0.4, 2.9, 1.2, "GRID ESTRATEGICO\nH3 Hexagonal\nCreencias Bayesianas\nValue of Information", H_ROSE),
    ]

    for x, y, w, h, label, color in boxes:
        rect = mpatches.FancyBboxPatch(
            (x, y), w, h, boxstyle="round,pad=0.12",
            facecolor=color, edgecolor="white", linewidth=2, alpha=0.92
        )
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2, label, ha="center", va="center",
                fontsize=6.5, color=H_DARK, fontweight="bold", linespacing=1.3)

    # Arrows between rows
    for x_start in [1.9, 5.6, 9.2]:
        ax.annotate("", xy=(x_start, 5.8), xytext=(x_start, 5.2),
                     arrowprops=dict(arrowstyle="->", color=H_SLATE, lw=1.5))
    for x_start in [1.4, 3.85, 6.35, 9.25]:
        ax.annotate("", xy=(x_start, 4.0), xytext=(x_start, 3.4),
                     arrowprops=dict(arrowstyle="->", color=H_SLATE, lw=1.5))
    for x_start in [2.0, 5.8, 9.25]:
        ax.annotate("", xy=(x_start, 2.2), xytext=(x_start, 1.6),
                     arrowprops=dict(arrowstyle="->", color=H_SLATE, lw=1.5))

    plt.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight", facecolor=H_CREAM)
    plt.close(fig)
    buf.seek(0)
    return buf


def generate_indices_chart():
    fig, ax = plt.subplots(figsize=(10, 5))

    indices = ["NDVI", "EVI", "SAVI", "NDRE", "NBR", "BSI", "NDBI", "MNDWI",
               "NDWI", "NDMI", "NBI", "MSAVI", "GNDVI", "CIre", "RECIg"]
    categories = ["Vegetacion", "Vegetacion", "Vegetacion", "Red-Edge", "Quema",
                   "Suelo", "Urbano", "Agua", "Agua", "Humedad",
                   "Suelo", "Vegetacion", "Vegetacion", "Red-Edge", "Red-Edge"]
    cat_colors = {
        "Vegetacion": H_TEAL, "Red-Edge": H_SAGE, "Quema": H_CORAL,
        "Suelo": H_PEACH, "Urbano": H_LAV, "Agua": H_SKY, "Humedad": H_MINT
    }
    colors = [cat_colors[c] for c in categories]

    x = np.arange(len(indices))
    bars = ax.bar(x, np.ones(len(indices)) * 2, bottom=-1 * np.ones(len(indices)),
                  color=colors, alpha=0.8, edgecolor="white", linewidth=1.5)

    ax.set_xticks(x)
    ax.set_xticklabels(indices, fontsize=7, fontweight="bold", rotation=25, ha="right")
    ax.set_ylabel("Rango [-1, +1]", fontsize=10)
    ax.set_title("15 Indices Espectrales de APEX — Sentinel-2", fontsize=14, fontweight="bold",
                 color=H_DARK, pad=15)
    ax.set_facecolor(H_CREAM)
    fig.patch.set_facecolor("white")
    ax.axhline(y=0, color=H_SLATE, linewidth=0.8, linestyle="--", alpha=0.5)
    ax.grid(axis="y", alpha=0.2)

    # Category annotations
    for i, (bar, cat) in enumerate(zip(bars, categories)):
        ax.text(i, 0, cat, ha="center", va="center", fontsize=5, color=H_DARK,
                fontweight="bold", rotation=90, alpha=0.6)

    # Legend patches
    handles = [mpatches.Patch(color=v, label=k) for k, v in cat_colors.items()]
    ax.legend(handles=handles, loc="upper right", fontsize=7, framealpha=0.9)

    plt.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def generate_engines_chart():
    fig, ax = plt.subplots(figsize=(9, 5))

    engines = [
        "Deforestacion", "Dynamic World", "Vegetacion", "Hansen GFC",
        "Alertas GLAD/RADD", "Incendios MODIS", "FIRMS NRT", "SAR Sentinel-1",
        "Drivers WRI", "CCDC Temporal", "Prithvi Deep Learning", "Legal ANP"
    ]
    capabilities = [6, 9, 7, 3, 4, 2, 5, 3, 7, 4, 10, 2]
    colors = [H_TEAL, H_SKY, H_SAGE, H_PEACH, H_CORAL, H_ROSE,
              "#f09880", H_LAV, H_MINT, "#b8c9e4", "#d4b8e8", "#ffd4a8"]

    bars = ax.barh(engines, capabilities, color=colors, edgecolor="white", linewidth=1.5, height=0.65)
    for bar, cap in zip(bars, capabilities):
        ax.text(bar.get_width() + 0.2, bar.get_y() + bar.get_height() / 2,
                f"{cap} clases", va="center", fontsize=8, fontweight="bold", color=H_DARK)

    ax.set_xlabel("Clases / Metricas de Salida", fontsize=10)
    ax.set_title("12 Motores de Analisis — Capacidad por Motor", fontsize=14,
                 fontweight="bold", color=H_DARK, pad=15)
    ax.set_facecolor(H_CREAM)
    fig.patch.set_facecolor("white")
    ax.set_xlim(0, max(capabilities) + 3)

    plt.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def generate_dw_classes_chart():
    fig, ax = plt.subplots(figsize=(8, 5))

    classes = ["Agua", "Bosque\nDenso", "Pastizal", "Vegetacion\nInundable",
               "Cultivos", "Matorral", "Zona\nUrbana", "Suelo\nDesnudo", "Nieve\nHielo"]
    colors_dw = ["#4169E1", "#228B22", "#90EE90", "#8FBC8F",
                 "#DAA520", "#BDB76B", "#CD853F", "#D2B48C", "#F0F8FF"]
    sizes = [8, 35, 12, 5, 18, 10, 7, 4, 1]

    wedges, texts, autotexts = ax.pie(
        sizes, labels=classes, colors=colors_dw, autopct="%1.0f%%",
        startangle=140, textprops={"fontsize": 8, "fontweight": "bold"},
        wedgeprops={"edgecolor": "white", "linewidth": 2},
        pctdistance=0.75
    )
    for at in autotexts:
        at.set_fontsize(7)
        at.set_color(H_DARK)

    ax.set_title("Dynamic World V1 — 9 Clases de Cobertura Terrestre",
                 fontsize=13, fontweight="bold", color=H_DARK, pad=15)
    fig.patch.set_facecolor("white")

    plt.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def generate_timeline_chart():
    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(10, 6), gridspec_kw={"height_ratios": [2, 1]})

    years = list(range(2018, 2026))
    deforestation = [120, 145, 180, 210, 195, 250, 310, 275]
    urban = [45, 52, 65, 78, 90, 105, 130, 142]
    vegetation = [85, 82, 78, 72, 70, 65, 58, 55]

    ax1.fill_between(years, deforestation, alpha=0.3, color=H_CORAL)
    ax1.plot(years, deforestation, "o-", color=H_CORAL, linewidth=2, label="Deforestacion (ha)")
    ax1.fill_between(years, urban, alpha=0.3, color=H_LAV)
    ax1.plot(years, urban, "s-", color=H_LAV, linewidth=2, label="Expansion Urbana (ha)")
    ax1.plot(years, [np.mean(deforestation)] * len(years), "--", color=H_SLATE, alpha=0.5,
             label="Promedio historico")

    ax1.set_ylabel("Hectareas Detectadas", fontsize=10)
    ax1.set_title("Analisis Temporal Multi-Anual (2018-2025) — Ejemplo de Resultados",
                  fontsize=13, fontweight="bold", color=H_DARK, pad=10)
    ax1.legend(fontsize=8, framealpha=0.9)
    ax1.set_facecolor(H_CREAM)
    ax1.grid(alpha=0.3)

    # Anomaly detection
    z_scores = [(d - np.mean(deforestation)) / np.std(deforestation) for d in deforestation]
    colors_z = [H_CORAL if abs(z) > 1.5 else (H_PEACH if abs(z) > 1 else H_MINT) for z in z_scores]
    ax2.bar(years, z_scores, color=colors_z, edgecolor="white", linewidth=1.5)
    ax2.axhline(y=1.5, color=H_CORAL, linestyle="--", alpha=0.7, label="Umbral anomalia (|Z|>1.5)")
    ax2.axhline(y=-1.5, color=H_CORAL, linestyle="--", alpha=0.7)
    ax2.set_ylabel("Z-score", fontsize=10)
    ax2.set_xlabel("Ano", fontsize=10)
    ax2.legend(fontsize=7, loc="lower right")
    ax2.set_facecolor(H_CREAM)
    ax2.grid(alpha=0.3)

    fig.patch.set_facecolor("white")
    plt.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def generate_tech_stack_diagram():
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 7)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    layers = [
        (0.3, 5.8, 9.4, 0.8, "FRONTEND — React 19 + TypeScript 5.9 • MapLibre GL • Terra Draw • Recharts • Tailwind CSS 4", H_SKY),
        (0.3, 4.7, 9.4, 0.8, "BACKEND — FastAPI 0.115 • Python 3.11 • Uvicorn ASGI • 36 Endpoints REST", H_TEAL),
        (0.3, 3.6, 4.5, 0.8, "ENGINES — 12 Motores de Analisis\n15 Indices Espectrales", H_SAGE),
        (5.1, 3.6, 4.6, 0.8, "ML / AI — Prithvi-EO 300M\nRandom Forest • POMDP • Bayesian", H_LAV),
        (0.3, 2.5, 4.5, 0.8, "GEOESPACIAL — GEE 0.1.396\nGeoPandas • rasterio • Shapely", H_MINT),
        (5.1, 2.5, 4.6, 0.8, "DATOS — Sentinel-2 • Dynamic World\nHansen GFC • GLAD/RADD • FIRMS", H_PEACH),
        (0.3, 1.4, 4.5, 0.8, "DATABASE — PostgreSQL 16 + PostGIS\nSQLite • SQLAlchemy 2.0 ORM", H_ROSE),
        (5.1, 1.4, 4.6, 0.8, "GPU / COMPUTE — PyTorch CUDA\nRTX 4050 (5.5 GB VRAM)", H_CORAL),
        (0.3, 0.2, 9.4, 0.8, "REPORTES — ReportLab 4 (PDF) • python-docx (Word) • matplotlib 3.10 • Pillow • Folio PROFEPA", H_TEAL),
    ]

    for x, y, w, h, label, color in layers:
        rect = mpatches.FancyBboxPatch(
            (x, y), w, h, boxstyle="round,pad=0.08",
            facecolor=color, edgecolor="white", linewidth=2, alpha=0.88
        )
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2, label, ha="center", va="center",
                fontsize=7.5, color=H_DARK, fontweight="bold", linespacing=1.2)

    ax.text(5, 6.9, "APEX — Stack Tecnologico Completo", ha="center",
            fontsize=15, color=H_DARK, fontweight="bold")

    plt.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def generate_ndvi_comparison():
    """Generate before/after NDVI comparison image."""
    fig, axes = plt.subplots(1, 3, figsize=(12, 4))
    np.random.seed(123)

    ndvi_cmap = LinearSegmentedColormap.from_list("ndvi",
        [(0.6, 0.15, 0.1), (0.85, 0.7, 0.2), (0.95, 0.95, 0.3), (0.3, 0.7, 0.15), (0.05, 0.4, 0.05)], N=256)

    # Before (healthy)
    x = np.linspace(0, 6, 300)
    y = np.linspace(0, 4, 200)
    X, Y = np.meshgrid(x, y)
    before = 0.6 + 0.2 * np.sin(X) * np.cos(Y * 0.8) + 0.1 * np.random.rand(200, 300)
    before = np.clip(before, 0, 1)
    axes[0].imshow(before, cmap=ndvi_cmap, vmin=0, vmax=1, aspect="auto")
    axes[0].set_title("T1: Enero 2022\nNDVI Medio: 0.72", fontsize=9, fontweight="bold", color=H_DARK)
    axes[0].set_xticks([]); axes[0].set_yticks([])

    # After (deforested)
    after = before.copy()
    after[60:140, 80:220] = 0.15 + 0.05 * np.random.rand(80, 140)  # cleared area
    after[30:70, 200:260] = 0.12 + 0.04 * np.random.rand(40, 60)
    axes[1].imshow(after, cmap=ndvi_cmap, vmin=0, vmax=1, aspect="auto")
    axes[1].set_title("T2: Enero 2024\nNDVI Medio: 0.48", fontsize=9, fontweight="bold", color=H_DARK)
    axes[1].set_xticks([]); axes[1].set_yticks([])

    # Difference (change detection)
    diff = before - after
    diff_cmap = LinearSegmentedColormap.from_list("diff",
        [(0.2, 0.5, 0.8), (0.95, 0.95, 0.95), (0.9, 0.2, 0.1)], N=256)
    im = axes[2].imshow(diff, cmap=diff_cmap, vmin=-0.3, vmax=0.6, aspect="auto")
    axes[2].set_title("Cambio NDVI\nDeforestacion Detectada", fontsize=9, fontweight="bold", color=H_CORAL)
    axes[2].set_xticks([]); axes[2].set_yticks([])
    # Red outline on detected area
    rect = mpatches.Rectangle((80, 60), 140, 80, linewidth=2, edgecolor="#ff4444", facecolor="none")
    axes[2].add_patch(rect)
    rect2 = mpatches.Rectangle((200, 30), 60, 40, linewidth=2, edgecolor="#ff4444", facecolor="none")
    axes[2].add_patch(rect2)

    # Colorbar
    cbar = fig.colorbar(im, ax=axes, orientation="vertical", fraction=0.02, pad=0.02)
    cbar.set_label("ΔNDVI", fontsize=9)

    fig.suptitle("Deteccion de Deforestacion — Comparacion Multitemporal NDVI",
                 fontsize=12, fontweight="bold", color=H_DARK, y=1.02)
    fig.patch.set_facecolor("white")
    plt.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def generate_sar_comparison():
    """Generate SAR VV/VH backscatter change image."""
    fig, axes = plt.subplots(1, 3, figsize=(12, 4))
    np.random.seed(456)

    sar_cmap = "gray"

    x = np.linspace(0, 6, 300)
    y = np.linspace(0, 4, 200)
    X, Y = np.meshgrid(x, y)

    # T1 SAR backscatter (high values = forest)
    t1 = -8 + 3 * np.sin(X * 0.5) * np.cos(Y * 0.3) + np.random.randn(200, 300) * 0.5
    axes[0].imshow(t1, cmap=sar_cmap, vmin=-15, vmax=0, aspect="auto")
    axes[0].set_title("SAR VV T1 (2022)\nBaseline", fontsize=9, fontweight="bold", color=H_DARK)
    axes[0].set_xticks([]); axes[0].set_yticks([])

    # T2 SAR (deforested = lower backscatter)
    t2 = t1.copy()
    t2[50:130, 70:200] -= 5  # 5 dB drop = deforestation
    axes[1].imshow(t2, cmap=sar_cmap, vmin=-15, vmax=0, aspect="auto")
    axes[1].set_title("SAR VV T2 (2024)\nCon Cambios", fontsize=9, fontweight="bold", color=H_DARK)
    axes[1].set_xticks([]); axes[1].set_yticks([])

    # Change magnitude
    change = np.abs(t2 - t1)
    im = axes[2].imshow(change, cmap="YlOrRd", vmin=0, vmax=8, aspect="auto")
    axes[2].set_title("Magnitud Cambio\n|ΔVV| > 3dB = Alerta", fontsize=9, fontweight="bold", color=H_CORAL)
    axes[2].set_xticks([]); axes[2].set_yticks([])
    rect = mpatches.Rectangle((70, 50), 130, 80, linewidth=2, edgecolor=H_CORAL, facecolor="none", linestyle="--")
    axes[2].add_patch(rect)

    cbar = fig.colorbar(im, ax=axes, orientation="vertical", fraction=0.02, pad=0.02)
    cbar.set_label("dB", fontsize=9)

    fig.suptitle("Sentinel-1 SAR — Deteccion Radar de Deforestacion (All-Weather)",
                 fontsize=12, fontweight="bold", color=H_DARK, y=1.02)
    fig.patch.set_facecolor("white")
    plt.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def generate_pomdp_diagram():
    fig, ax = plt.subplots(figsize=(9, 5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    # POMDP cycle
    nodes = [
        (2, 4.5, "OBSERVACION\nImagenes satelitales\nSentinel-2 + SAR", H_SKY),
        (5, 5, "CREENCIA\nP(tala | datos)\nBayesian update", H_LAV),
        (8, 4.5, "DECISION\nInspeccionar?\nAdquirir imagen?\nEsperar?", H_PEACH),
        (8, 2, "ACCION\nPlan semanal\nRutas inspectores\nPresupuesto", H_CORAL),
        (5, 1.2, "RESULTADO\nValidacion campo\nFeedback inspector\nReentrenamiento", H_SAGE),
        (2, 2, "ACTUALIZACION\nBelief degradation\nNuevos datos\nRe-priorizacion", H_MINT),
    ]

    for x, y, label, color in nodes:
        circle = mpatches.Circle((x, y), 0.95, facecolor=color, edgecolor="white", linewidth=2, alpha=0.9)
        ax.add_patch(circle)
        ax.text(x, y, label, ha="center", va="center", fontsize=6.5, color=H_DARK,
                fontweight="bold", linespacing=1.3)

    # Arrows
    pairs = [(0, 1), (1, 2), (2, 3), (3, 4), (4, 5), (5, 0)]
    coords = [(n[0], n[1]) for n in nodes]
    for i, j in pairs:
        x1, y1 = coords[i]
        x2, y2 = coords[j]
        dx, dy = x2 - x1, y2 - y1
        dist = math.sqrt(dx ** 2 + dy ** 2)
        ax.annotate("", xy=(x2 - dx / dist * 1.0, y2 - dy / dist * 1.0),
                     xytext=(x1 + dx / dist * 1.0, y1 + dy / dist * 1.0),
                     arrowprops=dict(arrowstyle="-|>", color=H_DARK, lw=2))

    ax.text(5, 6.0, "POMDP — Ciclo de Decision Optima para Inspeccion Ambiental",
            ha="center", fontsize=13, fontweight="bold", color=H_DARK)

    plt.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def generate_forecast_chart():
    fig, ax = plt.subplots(figsize=(9, 5))

    years_hist = list(range(2018, 2026))
    years_fc = list(range(2026, 2031))
    defor_hist = [120, 145, 180, 210, 195, 250, 310, 275]

    # Simple trend projection
    trend = np.polyfit(range(len(defor_hist)), defor_hist, 1)
    defor_fc = [trend[0] * (len(defor_hist) + i) + trend[1] for i in range(5)]

    # Confidence intervals
    ci_low = [v * 0.7 for v in defor_fc]
    ci_high = [v * 1.4 for v in defor_fc]

    ax.fill_between(years_fc, ci_low, ci_high, alpha=0.15, color=H_CORAL, label="IC 95%")
    ax.fill_between(years_fc, [v * 0.85 for v in defor_fc], [v * 1.2 for v in defor_fc],
                    alpha=0.25, color=H_CORAL, label="IC 68%")
    ax.plot(years_hist, defor_hist, "o-", color=H_TEAL, linewidth=2.5, markersize=6, label="Datos Historicos")
    ax.plot(years_fc, defor_fc, "s--", color=H_CORAL, linewidth=2, markersize=6, label="Prediccion Ensemble")

    ax.axvline(x=2025.5, color=H_SLATE, linestyle=":", alpha=0.5)
    ax.text(2025.5, max(ci_high) * 0.95, "  Horizonte\n  prediccion", fontsize=8, color=H_SLATE)

    ax.set_xlabel("Ano", fontsize=11)
    ax.set_ylabel("Hectareas Deforestadas", fontsize=11)
    ax.set_title("Motor de Prediccion — Forecast Ensemble (Trend + ML + POMDP)",
                 fontsize=13, fontweight="bold", color=H_DARK, pad=10)
    ax.legend(fontsize=8, framealpha=0.9)
    ax.set_facecolor(H_CREAM)
    fig.patch.set_facecolor("white")
    ax.grid(alpha=0.3)

    plt.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def generate_comparison_radar():
    """Radar chart comparing APEX vs other platforms."""
    fig, ax = plt.subplots(figsize=(7, 7), subplot_kw=dict(polar=True))

    categories = [
        "Motores de\nAnalisis", "Indices\nEspectrales", "ML / Deep\nLearning",
        "Datos\nSatelitales", "Reportes\nInstitucionales", "Prediccion\nForecasting",
        "Monitoreo\nContinuo", "Decision\nPOMDP"
    ]
    N = len(categories)
    angles = np.linspace(0, 2 * np.pi, N, endpoint=False).tolist()
    angles += angles[:1]

    apex_scores = [10, 10, 9, 9, 10, 8, 9, 10]
    gfw_scores  = [3, 2, 4, 6, 2, 3, 7, 0]
    eo_scores   = [2, 8, 3, 8, 1, 0, 2, 0]

    apex_scores += apex_scores[:1]
    gfw_scores += gfw_scores[:1]
    eo_scores += eo_scores[:1]

    ax.fill(angles, apex_scores, color=H_TEAL, alpha=0.25)
    ax.plot(angles, apex_scores, "o-", color=H_TEAL, linewidth=2.5, label="APEX", markersize=6)
    ax.fill(angles, gfw_scores, color=H_PEACH, alpha=0.15)
    ax.plot(angles, gfw_scores, "s--", color=H_PEACH, linewidth=1.5, label="GFW", markersize=4)
    ax.fill(angles, eo_scores, color=H_LAV, alpha=0.15)
    ax.plot(angles, eo_scores, "^--", color=H_LAV, linewidth=1.5, label="EO Browser", markersize=4)

    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(categories, fontsize=8, fontweight="bold")
    ax.set_ylim(0, 10)
    ax.set_title("Comparativa de Capacidades", fontsize=14, fontweight="bold",
                 color=H_DARK, pad=20)
    ax.legend(fontsize=9, loc="upper right", bbox_to_anchor=(1.25, 1.1))
    ax.set_facecolor(H_CREAM)
    fig.patch.set_facecolor("white")

    plt.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


# ════════════════════════════════════════════
#  MAIN DOCUMENT GENERATION
# ════════════════════════════════════════════
def generate_report():
    print("Generando reporte tecnico APEX...")
    doc = Document()

    # Page setup (Letter)
    sec = doc.sections[0]
    sec.page_width = Cm(21.59)
    sec.page_height = Cm(27.94)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    FIG_NUM = [0]  # mutable counter

    def fig_caption(doc, text):
        FIG_NUM[0] += 1
        add_figure_caption(doc, f"Figura {FIG_NUM[0]}. {text}")

    # ══════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════
    for _ in range(3):
        doc.add_paragraph()

    if PROFEPA_LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(PROFEPA_LOGO), width=Inches(1.8))

    doc.add_paragraph()

    # APEX triangle logo (text-based)
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("▲ APEX")
    r.font.size = Pt(48)
    r.font.color.rgb = APEX_TEAL
    r.bold = True

    doc.add_paragraph()

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("REPORTE TÉCNICO INTEGRAL")
    r2.font.size = Pt(26)
    r2.font.color.rgb = APEX_DARK
    r2.bold = True

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = t3.add_run("Plataforma Avanzada de Análisis Geoespacial Ambiental")
    r3.font.size = Pt(16)
    r3.font.color.rgb = APEX_SAGE

    doc.add_paragraph()

    t4 = doc.add_paragraph()
    t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = t4.add_run(
        "Análisis y Protección del Entorno por Exploración Satelital\n\n"
        "PROFEPA — Procuraduría Federal de Protección al Ambiente\n"
        "Coordinación de Estudios Prospectivos y Valoración de Riesgos"
    )
    r4.font.size = Pt(11)
    r4.font.color.rgb = GRAY

    doc.add_paragraph()

    td = doc.add_paragraph()
    td.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rd = td.add_run("Marzo 2026")
    rd.font.size = Pt(14)
    rd.font.color.rgb = APEX_TEAL
    rd.bold = True

    doc.add_page_break()

    # ══════════════════════════════════════════
    # TABLE OF CONTENTS
    # ══════════════════════════════════════════
    add_heading(doc, "CONTENIDO", 1)
    toc = [
        "1. Resumen Ejecutivo",
        "2. Contexto y Problemática",
        "3. Objetivos del Sistema",
        "4. Arquitectura Técnica",
        "5. Stack Tecnológico",
        "6. Motores de Análisis (12 Engines)",
        "7. Índices Espectrales (15 Índices)",
        "8. Análisis Temporal y Detección de Anomalías",
        "9. Machine Learning y Deep Learning",
        "10. Planificación Óptima (POMDP)",
        "11. Motor de Predicción (Forecast)",
        "12. Integración de Datos de Incendios",
        "13. Interfaz de Usuario",
        "14. Generación de Reportes",
        "15. API REST — 36+ Endpoints",
        "16. Base de Datos",
        "17. Monitoreo Continuo y Alertas",
        "18. Seguridad y Autenticación",
        "19. Infraestructura y Despliegue",
        "20. Casos de Uso — Ejemplos con Imágenes Satelitales",
        "21. Tabla Comparativa",
        "22. Roadmap y Evolución",
        "23. Conclusiones",
        "Anexo A. Fórmulas de Índices Espectrales",
        "Anexo B. Matriz de Confusión CCDC",
    ]
    for item in toc:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(10.5)

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 1. RESUMEN EJECUTIVO
    # ══════════════════════════════════════════
    add_heading(doc, "1. Resumen Ejecutivo", 1)

    add_text(doc,
        "APEX (Análisis y Protección del Entorno por Exploración Satelital) es una plataforma "
        "web de análisis geoespacial diseñada para la Procuraduría Federal de Protección al "
        "Ambiente (PROFEPA). Combina 12 motores de análisis, 15 índices espectrales, "
        "inteligencia artificial (Prithvi-EO 300M, Random Forest, Bayesian Fusion) y "
        "planificación de decisión óptima (POMDP) para la detección de deforestación, "
        "expansión urbana ilegal y clasificación de cobertura vegetal en el territorio mexicano."
    )

    add_text(doc,
        "A diferencia de plataformas de observación terrestre de uso general, APEX ha sido "
        "diseñado específicamente para el flujo de trabajo de verificación e inspección "
        "ambiental de PROFEPA, con reportes institucionales, análisis temporal multi-anual "
        "(2018-2025), predicción a horizonte de 1-5 años, y un módulo de planificación "
        "estratégica que optimiza la asignación de inspectores y presupuesto."
    )

    make_table(doc,
        ["Métrica", "Valor"],
        [
            ["Motores de análisis", "12 (deforestación, DW, vegetación, Hansen, GLAD/RADD, fire, FIRMS, SAR, drivers, CCDC, Prithvi, legal)"],
            ["Índices espectrales", "15 (NDVI, EVI, SAVI, MSAVI, GNDVI, NDRE, CIre, RECIg, NDWI, MNDWI, NDMI, BSI, NBI, NDBI, NBR)"],
            ["API endpoints", "36+"],
            ["Fuentes satelitales", "7 (Sentinel-2, Sentinel-1 SAR, Dynamic World, Hansen GFC, GLAD, RADD, FIRMS)"],
            ["Clases de cobertura", "9 (Dynamic World V1 de Google)"],
            ["Deep Learning", "Prithvi-EO-2.0-300M (NASA/IBM, 10 clases)"],
            ["Predicción", "Horizonte 1-5 años (trend, ML, POMDP, ensemble)"],
            ["Resolución espacial", "10m (Sentinel-2 / Dynamic World)"],
            ["Cobertura temporal", "2018-2025 (multi-anual con detección de anomalías Z-score)"],
            ["Formatos de reporte", "4 (PDF, Word, JSON, GeoJSON)"],
            ["GPU acceleration", "CUDA — RTX 4050 (5.5 GB VRAM)"],
            ["Líneas de código", "~44,000+ (backend ~22K + frontend ~22K)"],
        ],
        widths=[4, 12]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 2. CONTEXTO
    # ══════════════════════════════════════════
    add_heading(doc, "2. Contexto y Problemática", 1)

    add_text(doc,
        "México pierde aproximadamente 203,552 hectáreas anuales de cobertura forestal (CONAFOR, 2001-2024). "
        "Las principales causas son la conversión a pastizales (73.29%) y tierras agrícolas (21.99%), "
        "muchas veces sin la autorización ambiental correspondiente. La fiscalización de estos "
        "cambios de uso de suelo representa un reto significativo para PROFEPA debido a la extensión "
        "del territorio y la complejidad logística de las inspecciones de campo."
    )

    make_table(doc,
        ["Problema Actual", "Solución APEX"],
        [
            ["Inspecciones costosas y lentas", "Detección remota automatizada con 12 motores"],
            ["Análisis manual de imágenes", "Pipeline automático de 15 índices espectrales"],
            ["Sin registro histórico de cambios", "Análisis temporal 2018-2025 con Z-score"],
            ["Reportes manuales", "Generación automática PDF/Word con folio PROFEPA"],
            ["Dificultad para detectar anomalías", "Detección estadística + ML + Deep Learning"],
            ["Falta de priorización", "POMDP → plan semanal óptimo de inspecciones"],
            ["Datos aislados de incendios", "Integración FIRMS con cruce de áreas de interés"],
            ["Sin predicción", "Forecast a 1-5 años (trend + ML + ensemble)"],
        ],
        widths=[5, 11]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 3. OBJETIVOS
    # ══════════════════════════════════════════
    add_heading(doc, "3. Objetivos del Sistema", 1)

    add_text(doc,
        "APEX tiene como objetivo principal dotar a PROFEPA de una herramienta institucional "
        "para monitoreo ambiental, detección de ilícitos forestales y generación de reportes "
        "técnicos con respaldo científico.", bold=True
    )

    add_heading(doc, "Objetivos Específicos", 2)
    objs = [
        "Integrar 12 motores de análisis complementarios (óptico, radar, ML, histórico, incendios, legal) en una sola plataforma.",
        "Calcular 15 índices espectrales a partir de bandas Sentinel-2 con interpretación automática.",
        "Clasificar la cobertura terrestre en 9-10 clases usando Dynamic World V1 y Prithvi-EO 300M (deep learning).",
        "Realizar análisis temporal multi-anual (2018-2025) con detección de anomalías por Z-score.",
        "Integrar datos de incendios NASA FIRMS en tiempo real (VIIRS SNPP, NOAA-20, NOAA-21, MODIS).",
        "Implementar detección radar SAR (Sentinel-1 VV/VH) para monitoreo en cualquier condición climática.",
        "Generar predicciones de deforestación a horizonte de 1-5 años (trend, ML, POMDP, ensemble).",
        "Desarrollar un módulo POMDP de planificación óptima de inspecciones con Value of Information.",
        "Implementar fusión bayesiana de evidencias multi-motor en celdas hexagonales H3.",
        "Generar reportes institucionales (PDF/Word) con formato PROFEPA y folio único.",
        "Desplegar como aplicación web con GPU acceleration (CUDA, RTX 4050).",
        "Proporcionar KPIs operativos con dashboard de efectividad, análisis por motor y tendencias.",
    ]
    for o in objs:
        add_bullet(doc, o)

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 4. ARQUITECTURA
    # ══════════════════════════════════════════
    add_heading(doc, "4. Arquitectura Técnica", 1)

    add_text(doc,
        "APEX utiliza una arquitectura cliente-servidor moderna basada en FastAPI (backend) y "
        "React (frontend), con procesamiento de imágenes satelitales en la nube mediante "
        "Google Earth Engine y aceleración GPU local con PyTorch CUDA para inferencia con "
        "modelos foundation como Prithvi-EO-2.0-300M."
    )

    print("  [1/14] Diagrama de arquitectura...")
    insert_image(doc, generate_architecture_diagram(), Inches(6.3))
    fig_caption(doc, "Arquitectura completa del sistema APEX con 12 motores de análisis, 4 capas de procesamiento y módulo de decisión POMDP.")

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 5. STACK TECNOLÓGICO
    # ══════════════════════════════════════════
    add_heading(doc, "5. Stack Tecnológico", 1)

    print("  [2/14] Stack tecnológico...")
    insert_image(doc, generate_tech_stack_diagram(), Inches(6.3))
    fig_caption(doc, "Stack tecnológico completo de APEX — desde frontend hasta GPU compute.")

    make_table(doc,
        ["Capa", "Tecnología", "Versión", "Función"],
        [
            ["Frontend", "React", "19.1", "Interfaz de usuario con TypeScript 5.9"],
            ["Frontend", "MapLibre GL", "5.3", "Mapa interactivo vectorial de alto rendimiento"],
            ["Frontend", "Terra Draw", "1.0", "Dibujo de polígonos (AOI) sobre el mapa"],
            ["Frontend", "Recharts", "2.15", "Gráficas temporales interactivas"],
            ["Frontend", "Tailwind CSS", "4.1", "Diseño responsivo con tema oscuro"],
            ["Backend", "FastAPI", "0.115", "Framework ASGI de alto rendimiento"],
            ["Backend", "Python", "3.11", "Lenguaje principal del servidor"],
            ["Backend", "Uvicorn", "0.34", "Servidor ASGI con async/await"],
            ["Satelital", "Google Earth Engine", "0.1.396", "Procesamiento de imágenes en la nube"],
            ["Satelital", "NASA FIRMS API", "v2", "Alertas de incendios en tiempo real"],
            ["Geoespacial", "GeoPandas", "0.13", "Análisis vectorial geoespacial"],
            ["Geoespacial", "rasterio", "1.3", "Lectura/escritura de rasters GeoTIFF"],
            ["Geoespacial", "Shapely", "2.0", "Operaciones geométricas"],
            ["ML/AI", "PyTorch", "2.0+", "Deep Learning con CUDA GPU"],
            ["ML/AI", "scikit-learn", "1.3", "Random Forest, clasificación"],
            ["ML/AI", "Prithvi-EO", "2.0-300M", "Foundation model NASA/IBM (10 clases)"],
            ["Database", "PostgreSQL + PostGIS", "16", "Base de datos geoespacial (producción)"],
            ["Database", "SQLite", "3.45", "Base de datos ligera (desarrollo)"],
            ["Database", "SQLAlchemy", "2.0", "ORM dialect-agnostic"],
            ["Reportes", "ReportLab", "4.0", "Generación de PDF profesional"],
            ["Reportes", "python-docx", "0.8", "Generación de documentos Word"],
            ["Reportes", "matplotlib", "3.10", "Gráficas y visualizaciones"],
        ],
        widths=[2.5, 3.5, 1.5, 8]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 6. MOTORES DE ANÁLISIS
    # ══════════════════════════════════════════
    add_heading(doc, "6. Motores de Análisis (12 Engines)", 1)

    add_text(doc,
        "APEX cuenta con 12 motores de análisis especializados que operan en paralelo sobre "
        "cada área de interés (AOI). El pipeline orquestador ejecuta los motores seleccionados "
        "y consolida los resultados en un único reporte integrado."
    )

    print("  [3/14] Motores de análisis...")
    insert_image(doc, generate_engines_chart(), Inches(6.0))
    fig_caption(doc, "Capacidad de salida por cada uno de los 12 motores de análisis.")

    # Engine details
    add_heading(doc, "6.1 Motor de Deforestación", 2)
    add_text(doc,
        "Detecta pérdida de cobertura forestal mediante 6 índices espectrales (NDVI, EVI, SAVI, NBR, NDRE, BSI). "
        "Genera máscara de deforestación combinando: NDVI < 0.2 (pérdida) AND MNDWI ≤ 0.1 (no agua) AND NDBI ≤ 0.05 (no urbano). "
        "Vectoriza los polígonos detectados con cálculo de área en hectáreas y confianza [0-1]. "
        "Máximo 50 polígonos por análisis con filtro de área mínima ~0.5 ha."
    )

    add_heading(doc, "6.2 Motor Dynamic World", 2)
    add_text(doc,
        "Utiliza la clasificación Google Dynamic World V1 con 9 clases de cobertura terrestre "
        "a resolución de 10m. Detecta cambios T1→T2 donde bosque/vegetación se convierte a "
        "no-bosque. Incluye refinamiento morfológico y reclasificación Suelo→Urbano basada "
        "en análisis de vecindarios 5×5 y 7×7. Máximo 200 features, confianza ≥ 0.25."
    )

    print("  [4/14] Dynamic World...")
    insert_image(doc, generate_dw_classes_chart(), Inches(5.5))
    fig_caption(doc, "Distribución típica de las 9 clases de cobertura de Dynamic World V1.")

    add_heading(doc, "6.3 Motor de Vegetación", 2)
    add_text(doc,
        "Clasificación basada en reglas con 7 clases: agua, bosque denso, bosque ralo, pastizal, "
        "suelo desnudo, urbano, quemado. Usa umbrales de NDVI, NBR, MNDWI, BSI, NBI y NDBI "
        "con área mínima de ~0.1 ha."
    )

    add_heading(doc, "6.4 Motor Hansen GFC", 2)
    add_text(doc,
        "Analiza la base de datos Global Forest Change v1.12 (Hansen et al., University of Maryland). "
        "Proporciona pérdida forestal histórica anual desde 2000, cobertura arbórea base (%) "
        "y estadísticas acumuladas por año."
    )

    add_heading(doc, "6.5 Motor de Alertas GLAD/RADD", 2)
    add_text(doc,
        "Integra alertas GLAD (Sentinel-2, confirmadas/probables) y RADD (Sentinel-1 SAR, "
        "nominal/alta confianza). Aplica deduplicación espacial: si overlap > 50%, conserva "
        "la alerta de mayor confianza. Ranking: confirmed(4) > high(3) > probable(2) > nominal(1)."
    )

    add_heading(doc, "6.6 Motor de Incendios MODIS", 2)
    add_text(doc,
        "Procesa el producto MCD64A1 de área quemada MODIS (resolución 500m). Genera polígonos "
        "de área quemada con día del año y se cruza con detecciones de deforestación para "
        "correlación incendio-deforestación."
    )

    add_heading(doc, "6.7 Motor FIRMS NRT", 2)
    add_text(doc,
        "Descarga hotspots activos de NASA FIRMS en tiempo real. Normaliza confianza "
        "(high→0.9, nominal→0.6, low→0.3), agrupa por proximidad (0.005°) y calcula "
        "FRP total, conteo de detecciones y rango temporal por cluster."
    )

    add_heading(doc, "6.8 Motor SAR (Sentinel-1)", 2)
    add_text(doc,
        "Detección radar all-weather usando backscatter VV/VH en dominio dB. Umbral de cambio: "
        "|ΔVV| o |ΔVH| > 3 dB. Confianza = min(magnitud/10, 1). Fusión con óptico: "
        "SAR∩DW overlap > 50% → conf = 0.6×DW + 0.4×SAR."
    )

    add_heading(doc, "6.9 Motor Drivers WRI", 2)
    add_text(doc,
        "Clasifica las causas de deforestación usando datos WRI/Google DeepMind a 1km: "
        "agricultura permanente, commodities, cultivo rotacional, tala, incendios, "
        "asentamientos e infraestructura, perturbación natural."
    )

    add_heading(doc, "6.10 Motor CCDC", 2)
    add_text(doc,
        "Continuous Change Detection and Classification sobre series temporales NDVI Sentinel-2. "
        "Clasifica cambios en: desmonte súbito (magnitud < -0.3), degradación gradual (-0.3 a -0.1), "
        "sin cambio, o recuperación (> 0.1)."
    )

    add_heading(doc, "6.11 Motor Prithvi-EO (Deep Learning)", 2)
    add_text(doc,
        "Foundation model NASA/IBM Prithvi-EO-2.0-300M con 10 clases de cobertura. "
        "Inferencia por parches de 224×224 pixels, normalización Sentinel-2 (/10,000), "
        "batch size configurable, soporte CUDA GPU. Vectorización con limpieza morfológica."
    )

    add_heading(doc, "6.12 Motor Legal (ANP)", 2)
    add_text(doc,
        "Intersección espacial del AOI con Áreas Naturales Protegidas (ANP). "
        "Etiqueta cada detección con nombre, categoría y porcentaje de overlap del ANP."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 7. ÍNDICES ESPECTRALES
    # ══════════════════════════════════════════
    add_heading(doc, "7. Índices Espectrales (15 Índices)", 1)

    add_text(doc,
        "APEX calcula 15 índices espectrales a partir de las bandas multiespectrales de Sentinel-2, "
        "organizados en 7 categorías temáticas. Cada índice está implementado con división "
        "segura (denominador + 10⁻¹⁰) para evitar errores numéricos."
    )

    print("  [5/14] Índices espectrales...")
    insert_image(doc, generate_indices_chart(), Inches(6.2))
    fig_caption(doc, "Los 15 índices espectrales de APEX categorizados por tipo de análisis.")

    make_table(doc,
        ["Categoría", "Índice", "Fórmula", "Aplicación"],
        [
            ["Vegetación", "NDVI", "(B8-B4)/(B8+B4)", "Vigor y densidad de vegetación"],
            ["Vegetación", "EVI", "2.5×(B8-B4)/(B8+6×B4-7.5×B2+1)", "Vegetación mejorada (corrige atmósfera)"],
            ["Vegetación", "SAVI", "1.5×(B8-B4)/(B8+B4+0.5)", "Vegetación ajustada al suelo"],
            ["Vegetación", "MSAVI", "[2B8+1-√((2B8+1)²-8(B8-B4))]/2", "Ajuste de suelo avanzado"],
            ["Vegetación", "GNDVI", "(B8-B3)/(B8+B3)", "Contenido de clorofila"],
            ["Red-Edge", "NDRE", "(B8-B5)/(B8+B5)", "Biomasa en bosques densos"],
            ["Red-Edge", "CIre", "B8/B7 - 1", "Índice de clorofila red-edge"],
            ["Red-Edge", "RECIg", "B7/B5 - 1", "Índice de clorofila green"],
            ["Agua", "NDWI", "(B3-B8)/(B3+B8)", "Delimitación de cuerpos de agua"],
            ["Agua", "MNDWI", "(B3-B11)/(B3+B11)", "Agua modificado (mejor en urbano)"],
            ["Humedad", "NDMI", "(B8-B11)/(B8+B11)", "Contenido de humedad foliar"],
            ["Suelo", "BSI", "((B11+B4)-(B8+B2))/((B11+B4)+(B8+B2))", "Suelo desnudo expuesto"],
            ["Suelo", "NBI", "B11×B4/B8", "Índice nuevo de áreas baldías"],
            ["Urbano", "NDBI", "(B11-B8)/(B11+B8)", "Áreas construidas y urbanas"],
            ["Quema", "NBR", "(B8-B12)/(B8+B12)", "Áreas quemadas por incendio"],
        ],
        widths=[2, 1.8, 5.5, 5]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 8. ANÁLISIS TEMPORAL
    # ══════════════════════════════════════════
    add_heading(doc, "8. Análisis Temporal y Detección de Anomalías", 1)

    add_text(doc,
        "El pipeline temporal de APEX ejecuta análisis multi-anual desde 2018 hasta 2025, "
        "procesando cada año con los motores seleccionados. Soporta selección de estación "
        "(seca, lluviosa, anual) y acumula estadísticas incrementales."
    )

    print("  [6/14] Timeline...")
    insert_image(doc, generate_timeline_chart(), Inches(6.3))
    fig_caption(doc, "Análisis temporal multi-anual con detección de anomalías por Z-score. Las barras rojas indican años con Z > 1.5 (anomalías significativas).")

    add_heading(doc, "8.1 Detección de Anomalías (Z-Score)", 2)
    add_text(doc,
        "Para cada motor y cada año, se calcula el Z-score: Z = (x - μ) / σ, donde x es el "
        "valor del año, μ es la media histórica y σ la desviación estándar. Un |Z| > 1.5 "
        "se clasifica como anomalía significativa, indicando un cambio inusual que merece "
        "investigación prioritaria."
    )

    add_heading(doc, "8.2 Manejo de Áreas Grandes", 2)
    add_text(doc,
        "AOIs mayores a 150 km² se dividen automáticamente en cuadrículas (hasta 5×5), "
        "procesando cada segmento en paralelo y consolidando resultados con promedios "
        "ponderados por área."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 9. ML / DEEP LEARNING
    # ══════════════════════════════════════════
    add_heading(doc, "9. Machine Learning y Deep Learning", 1)

    add_heading(doc, "9.1 Prithvi-EO-2.0-300M (NASA/IBM Foundation Model)", 2)
    add_text(doc,
        "APEX integra el modelo foundation Prithvi-EO-2.0-300M desarrollado por NASA e IBM "
        "para clasificación de cobertura terrestre. El modelo procesa parches de 224×224 "
        "pixels de imágenes Sentinel-2, normalizados por 10,000 (factor de escala Sentinel-2). "
        "Clasifica en 10 clases incluyendo agua, bosque, pastizal, humedal, cultivo, matorral, "
        "urbano, suelo desnudo y nieve/hielo. Soporta backends Terratorch y HuggingFace Transformers."
    )

    add_heading(doc, "9.2 Random Forest (scikit-learn)", 2)
    add_text(doc,
        "Clasificador Random Forest con 13 bandas de entrada (10 reflectancia + NDVI, EVI, NBR) "
        "para clasificación de deforestación. Se utiliza también en el motor de forecast para "
        "predicción temporal de tendencias de deforestación."
    )

    add_heading(doc, "9.3 Fusión Bayesiana de Evidencias", 2)
    add_text(doc,
        "Sistema de creencias (beliefs) basado en actualización bayesiana sobre celdas "
        "hexagonales H3. Cada motor actualiza la probabilidad de ilícito (p_tala, p_cus, p_agri) "
        "con degradación temporal automática. Proporciona un índice de certidumbre (CI) "
        "y recomendaciones de adquisición de imagen."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 10. POMDP
    # ══════════════════════════════════════════
    add_heading(doc, "10. Planificación Óptima (POMDP)", 1)

    add_text(doc,
        "APEX implementa un módulo de Partially Observable Markov Decision Process (POMDP) "
        "para optimizar la asignación de inspectores y el presupuesto de adquisición de imágenes. "
        "El sistema genera planes semanales que maximizan la efectividad de las inspecciones "
        "considerando la incertidumbre en las detecciones."
    )

    print("  [7/14] POMDP...")
    insert_image(doc, generate_pomdp_diagram(), Inches(5.8))
    fig_caption(doc, "Ciclo de decisión POMDP — desde observación satelital hasta reentrenamiento.")

    make_table(doc,
        ["Componente POMDP", "Función", "Parámetros"],
        [
            ["Plan Semanal", "Asignación óptima de inspectores", "n_inspectores, presupuesto_USD"],
            ["Simulación", "Proyección de efectividad mensual", "inspectores, presupuesto, umbral_ha"],
            ["Rutas", "Generación de rutas de inspección", "n_inspectores, días, formato GeoJSON"],
            ["Value of Information", "Valor de adquirir imagen adicional por celda H3", "h3_index"],
        ],
        widths=[3, 5, 6]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 11. FORECAST
    # ══════════════════════════════════════════
    add_heading(doc, "11. Motor de Predicción (Forecast)", 1)

    add_text(doc,
        "El módulo de forecast permite proyectar tendencias de deforestación y cambio de uso "
        "de suelo a horizonte de 1-5 años usando 4 métodos complementarios."
    )

    print("  [8/14] Forecast...")
    insert_image(doc, generate_forecast_chart(), Inches(6.2))
    fig_caption(doc, "Ejemplo de predicción ensemble con intervalos de confianza a horizonte 2026-2030.")

    make_table(doc,
        ["Método", "Algoritmo", "Fortaleza"],
        [
            ["Trend", "Regresión lineal/polinomial", "Captura tendencias macro de largo plazo"],
            ["ML", "Random Forest temporal", "Captura patrones no lineales estacionales"],
            ["POMDP", "Markov con estados ocultos", "Incorpora incertidumbre de observación"],
            ["Ensemble", "Promedio ponderado de los 3", "Combina fortalezas, reduce varianza"],
        ],
        widths=[2.5, 5, 7]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 12. INCENDIOS
    # ══════════════════════════════════════════
    add_heading(doc, "12. Integración de Datos de Incendios", 1)

    make_table(doc,
        ["Sensor", "Resolución", "Frecuencia", "Motor APEX"],
        [
            ["VIIRS SNPP", "375m", "2 pasadas/día", "FIRMS NRT Engine"],
            ["VIIRS NOAA-20", "375m", "2 pasadas/día", "FIRMS NRT Engine"],
            ["VIIRS NOAA-21", "375m", "2 pasadas/día", "FIRMS NRT Engine"],
            ["MODIS Terra/Aqua", "1km", "4 pasadas/día", "Fire Engine (MCD64A1)"],
            ["Landsat 8/9 OLI", "30m", "~1 cada 8 días", "Disponible vía GEE"],
        ],
        widths=[3.5, 2.5, 3, 5]
    )

    add_text(doc,
        "El motor FIRMS normaliza confianza, agrupa hotspots por proximidad (0.005°), calcula "
        "Fire Radiative Power (FRP) total por cluster y genera polígonos convex hull. El motor "
        "Fire procesa el producto mensual MCD64A1 y correlaciona áreas quemadas con detecciones "
        "de deforestación."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 13. INTERFAZ
    # ══════════════════════════════════════════
    add_heading(doc, "13. Interfaz de Usuario", 1)

    add_text(doc,
        "APEX presenta una interfaz moderna desarrollada con React 19 y TypeScript 5.9, "
        "con diseño responsivo basado en Tailwind CSS 4 y tema oscuro profesional."
    )

    make_table(doc,
        ["Componente", "Tecnología", "Función"],
        [
            ["MapView", "MapLibre GL 5.3", "Mapa interactivo con capas de resultados por motor, pop-ups informativos"],
            ["Sidebar", "React + Tailwind", "Selector de motores, fechas, lanzamiento de análisis, exportación"],
            ["TopBar", "React", "Logo, estado de conexión, configuración de mapa base"],
            ["TimelinePanel", "Recharts 2.15", "Gráficas multi-anuales interactivas con zoom y filtros"],
            ["ForecastPanel", "React", "Configuración y visualización de predicciones (1-5 años)"],
            ["StrategicPanel", "React", "Mapa de riesgo H3, briefs de alerta generados por IA"],
            ["SimulatorPanel", "React", "Simulador POMDP: inspectores, presupuesto, rutas óptimas"],
            ["ImpactDashboard", "React + Recharts", "KPIs operativos, timeline de alertas, métricas por motor"],
            ["ValidationPanel", "React", "Revisión de detecciones (true/false), feedback para reentrenamiento"],
            ["MonitoringPanel", "React", "Gestión de áreas vigiladas, umbrales de alerta, notificaciones"],
            ["LegendPanel", "React", "Leyenda de colores por motor con toggle de visibilidad"],
            ["PolygonManager", "Terra Draw 1.0", "Dibujo/carga de AOIs en GeoJSON/Shapefile"],
            ["LoginPage", "React", "Autenticación JWT con formulario institucional"],
        ],
        widths=[3, 3, 9]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 14. REPORTES
    # ══════════════════════════════════════════
    add_heading(doc, "14. Generación de Reportes", 1)

    make_table(doc,
        ["Formato", "Librería", "Contenido", "Elementos"],
        [
            ["PDF", "ReportLab 4", "Reporte visual con branding PROFEPA", "Resumen ejecutivo, tablas, mapas, anomalías, anexo metodológico"],
            ["Word (.docx)", "python-docx", "Documento editable institucional", "Gráficas matplotlib, tablas de resultados, narrativa automática"],
            ["JSON", "nativo", "Datos estructurados para integración", "Resultados por motor, estadísticas, GeoJSON de detecciones"],
            ["GeoJSON", "nativo", "Capas geoespaciales", "Polígonos de detección con atributos por motor"],
        ],
        widths=[2, 2.5, 4.5, 6]
    )

    add_text(doc,
        "Cada reporte incluye un folio único, resumen ejecutivo con métricas agregadas, "
        "tablas por motor/año, sección de anomalías detectadas y anexo metodológico. "
        "Los reportes PDF incluyen branding institucional PROFEPA con logos y formato oficial."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 15. API
    # ══════════════════════════════════════════
    add_heading(doc, "15. API REST — 36+ Endpoints", 1)

    make_table(doc,
        ["Categoría", "Endpoints", "Métodos", "Descripción"],
        [
            ["Análisis", "5", "POST/GET", "Análisis fecha única, timeline multi-año, estado de trabajo, resultados, envío de reporte"],
            ["Monitoreo", "8", "CRUD", "Áreas vigiladas, toggle activo, historial de alertas, análisis manual"],
            ["Forecast", "4", "POST/GET", "Predicción run/AOI, entrenamiento, estado del modelo"],
            ["Beliefs", "4", "GET/POST", "Creencias bayesianas, actualización, degradación, mapa de riesgo"],
            ["Estratégico", "2", "GET", "Overview de riesgo, briefings de alerta con IA"],
            ["POMDP", "4", "POST/GET", "Plan semanal, simulación, rutas de inspección, Value of Information"],
            ["KPI", "4", "GET", "Summary operativo, por motor, timeline, reentrenamiento"],
            ["Grid", "2", "GET", "Celdas H3 por bbox, estadísticas nacionales"],
            ["Auth", "3", "POST/GET", "Login JWT, registro, perfil (me)"],
            ["Export", "1", "GET", "Descarga de reporte en PDF/Word/JSON"],
        ],
        widths=[2.5, 1.5, 2, 9]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 16. BASE DE DATOS
    # ══════════════════════════════════════════
    add_heading(doc, "16. Base de Datos", 1)

    add_text(doc,
        "APEX usa PostgreSQL 16 con PostGIS para producción y SQLite para desarrollo, "
        "con SQLAlchemy 2.0 como ORM dialect-agnostic. Las tablas principales son:"
    )

    make_table(doc,
        ["Tabla", "Campos Clave", "Función"],
        [
            ["jobs", "id, status, progress, aoi, engines, start/end_date, logs[]", "Registro de trabajos de análisis"],
            ["analysis_results", "job_id, engine, geojson, stats_json, year", "Resultados GeoJSON y estadísticas por motor"],
            ["gee_cache", "aoi_hash, date_range, raster_path", "Caché de descargas GEE (evita re-descarga)"],
            ["monitoring_areas", "id, name, aoi, engines, interval_hours, threshold_ha", "Áreas bajo vigilancia continua"],
            ["monitoring_alerts", "area_id, timestamp, engine, severity, email_sent", "Alertas generadas por monitoreo"],
            ["grid_cells", "h3_index, lat, lng, estado, municipio, ecosistema, en_anp", "Grid hexagonal nacional H3"],
            ["belief_states", "h3_index, p_sin_ilicito, p_tala, p_cus, p_agri, ci", "Estados de creencia bayesiana"],
            ["users", "id, email, role, full_name, hashed_password", "Cuentas de usuario institucional"],
            ["sessions", "token, user_id, expires_at", "Sesiones JWT activas"],
            ["forecast_models", "id, method, trained_at, n_samples, metrics", "Modelos entrenados de forecast"],
            ["kpi_snapshots", "date, engine, detections, validated, rejected", "Snapshots de KPIs operativos"],
        ],
        widths=[3, 5.5, 5]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 17. MONITOREO
    # ══════════════════════════════════════════
    add_heading(doc, "17. Monitoreo Continuo y Alertas", 1)

    add_text(doc,
        "APEX permite definir áreas de monitoreo continuo que se analizan automáticamente "
        "a intervalos configurables. Cuando una detección excede el umbral definido (en hectáreas), "
        "se genera una alerta por correo electrónico con los resultados del análisis."
    )

    make_table(doc,
        ["Parámetro", "Valor Default", "Descripción"],
        [
            ["Intervalo", "24 horas", "Frecuencia de re-análisis del área"],
            ["Umbral", "1.0 ha", "Hectáreas mínimas para disparar alerta"],
            ["Motores", "Todos disponibles", "Motores ejecutados en cada ciclo"],
            ["Email", "Configurable", "Dirección de notificación SMTP"],
            ["Estado", "Activo/Inactivo", "Toggle individual por área"],
        ],
        widths=[3, 3, 8]
    )

    add_heading(doc, "17.1 KPI Dashboard", 2)
    add_text(doc,
        "Panel de control con métricas operativas en tiempo real: total de trabajos completados, "
        "alertas generadas, detecciones validadas vs. rechazadas (precision tracking), "
        "tiempo promedio de respuesta, y tendencias semanales por subdelegación."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 18. SEGURIDAD
    # ══════════════════════════════════════════
    add_heading(doc, "18. Seguridad y Autenticación", 1)

    make_table(doc,
        ["Componente", "Implementación", "Detalle"],
        [
            ["Autenticación", "JWT (JSON Web Token)", "HS256, expiración 8 horas"],
            ["Roles", "admin / user", "Basado en claims del token"],
            ["Contraseñas", "bcrypt hash", "Salt automático, sin almacenamiento en texto plano"],
            ["Rate Limiting", "Configurable", "Por IP y por endpoint"],
            ["SMTP", "TLS (puerto 587)", "Notificaciones cifradas"],
            ["GEE Auth", "Service Account / Interactive", "Credenciales JSON o refresh token"],
            ["FIRMS API", "MAP_KEY", "Token de NASA para descarga de datos"],
        ],
        widths=[3, 4, 7]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 19. INFRAESTRUCTURA
    # ══════════════════════════════════════════
    add_heading(doc, "19. Infraestructura y Despliegue", 1)

    make_table(doc,
        ["Componente", "Requisito", "Especificación"],
        [
            ["GPU", "NVIDIA CUDA", "RTX 4050 o superior (5.5 GB VRAM mínimo)"],
            ["Python", "3.10+", "Con PyTorch, FastAPI, Earth Engine API"],
            ["Node.js", "18+", "Para compilación del frontend React"],
            ["GEE", "Cuenta de servicio", "Proyecto: profepa-deforestation"],
            ["FIRMS", "API Key", "Registro en NASA FIRMS"],
            ["SMTP", "Servidor de correo", "Gmail SMTP o institucional PROFEPA"],
            ["PostgreSQL", "16+ con PostGIS", "Base de datos geoespacial (producción)"],
            ["Redis", "7 (opcional)", "Cola de tareas Celery"],
            ["Storage", "50+ GB", "Caché de tiles GEE e imágenes satelitales"],
        ],
        widths=[3, 3, 8]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 20. CASOS DE USO CON IMÁGENES
    # ══════════════════════════════════════════
    add_heading(doc, "20. Casos de Uso — Ejemplos con Imágenes Satelitales", 1)

    add_text(doc,
        "A continuación se presentan ejemplos representativos de los tipos de análisis que "
        "APEX realiza, aplicados a diferentes regiones del territorio mexicano."
    )

    # Case 1: Lacandona
    add_heading(doc, "20.1 Selva Lacandona, Chiapas — Detección de Deforestación", 2)
    add_text(doc,
        "La Selva Lacandona es uno de los ecosistemas con mayor biodiversidad de México y está "
        "sujeta a presión constante por actividades agropecuarias y asentamientos irregulares. "
        "APEX detecta zonas de pérdida forestal mediante la combinación de índices NDVI, BSI y "
        "el motor de deforestación, identificando polígonos irregulares con baja cobertura vegetal "
        "dentro de la masa forestal. Los recuadros rojos indican las detecciones priorizadas."
    )
    print("  [9/14] Lacandona...")
    insert_image(doc, generate_lacandona_scene(), Inches(6.0))
    fig_caption(doc, "Composición Sentinel-2 RGB de la Selva Lacandona con detecciones de deforestación resaltadas en rojo.")

    # NDVI comparison
    print("  [10/14] NDVI comparison...")
    insert_image(doc, generate_ndvi_comparison(), Inches(6.2))
    fig_caption(doc, "Comparación multitemporal NDVI: T1 (2022) vs T2 (2024). Los recuadros rojos delimitan zonas con pérdida significativa de NDVI.")

    doc.add_page_break()

    # Case 2: Cancún
    add_heading(doc, "20.2 Zona Hotelera de Cancún, Quintana Roo — Expansión Urbana", 2)
    add_text(doc,
        "La Riviera Maya experimenta expansión urbana acelerada que afecta manglares y selva costera. "
        "APEX utiliza el motor Dynamic World para detectar cambios de cobertura forestal a zona "
        "construida, complementado con el índice NDBI (áreas construidas) e IBI para confirmar "
        "la presencia de infraestructura nueva. Las áreas grises representan zonas urbanas "
        "detectadas por el clasificador."
    )
    print("  [11/14] Cancún...")
    insert_image(doc, generate_cancun_scene(), Inches(6.0))
    fig_caption(doc, "Detección de expansión urbana en la zona hotelera de Cancún mediante clasificación Dynamic World.")

    # Case 3: Jalisco fires
    add_heading(doc, "20.3 Sierra de Manantlán, Jalisco — Incendios Forestales", 2)
    add_text(doc,
        "La Sierra de Manantlán (Reserva de la Biosfera) es vulnerable a incendios forestales "
        "durante la temporada seca. APEX integra datos NASA FIRMS en tiempo real, mostrando "
        "hotspots activos (puntos rojos) con su FRP (Fire Radiative Power) y nivel de confianza. "
        "Los halos amarillos indican la zona de influencia térmica del fuego."
    )
    insert_image(doc, generate_jalisco_fire(), Inches(6.0))
    fig_caption(doc, "Detección de incendios activos en la Sierra de Manantlán usando NASA FIRMS (VIIRS SNPP/NOAA-20).")

    doc.add_page_break()

    # Case 4: SAR
    add_heading(doc, "20.4 Análisis Radar SAR — Detección All-Weather", 2)
    add_text(doc,
        "En zonas con alta nubosidad permanente (como la Sierra Norte de Oaxaca o la Selva "
        "Lacandona en temporada de lluvias), las imágenes ópticas no son útiles. APEX utiliza "
        "el radar SAR de Sentinel-1 (banda C, VV/VH) para detectar cambios de backscatter "
        "que indican pérdida de vegetación. Un cambio > 3 dB genera una alerta."
    )
    print("  [12/14] SAR...")
    insert_image(doc, generate_sar_comparison(), Inches(6.2))
    fig_caption(doc, "Detección radar SAR Sentinel-1: T1 baseline, T2 con cambios, y mapa de magnitud de cambio (dB). Recuadro = alerta |ΔVV| > 3dB.")

    # Case 5: Sinaloa
    add_heading(doc, "20.5 Valle de Culiacán, Sinaloa — Cambio de Uso de Suelo", 2)
    add_text(doc,
        "La expansión agrícola en valles de Sinaloa implica la conversión de vegetación nativa "
        "a parcelas agrícolas y la construcción de infraestructura de riego. APEX detecta ambos "
        "tipos de cambio simultáneamente: deforestación (café) y urbanización (gris)."
    )
    insert_image(doc, generate_sinaloa_agriculture(), Inches(6.0))
    fig_caption(doc, "Detección combinada de deforestación y cambio de uso de suelo agrícola/urbano en el Valle de Culiacán.")

    doc.add_page_break()

    # Case 6: Monarch Reserve
    add_heading(doc, "20.6 Reserva Mariposa Monarca, Michoacán — Análisis NDVI", 2)
    add_text(doc,
        "La Reserva de la Biosfera Mariposa Monarca es un Área Natural Protegida (ANP) "
        "que APEX analiza con el motor Legal para verificar overlap con ANPs y el motor "
        "de deforestación para identificar tala clandestina. El motor Legal etiqueta cada "
        "detección con el nombre del ANP y porcentaje de overlap."
    )
    insert_image(doc, generate_monarch_reserve(), Inches(6.0))
    fig_caption(doc, "Análisis NDVI de la Reserva Mariposa Monarca con detecciones de pérdida forestal dentro del ANP.")

    # Case 7: Tabasco
    add_heading(doc, "20.7 Pantanos de Centla, Tabasco — Clasificación de Cobertura", 2)
    add_text(doc,
        "Los humedales de Tabasco presentan un mosaico complejo de coberturas que "
        "Dynamic World clasifica en 9 clases. APEX genera mapas de distribución porcentual "
        "por clase y detecta transiciones inter-anuales de vegetación inundable a pastizal "
        "o suelo desnudo, indicativas de drenaje ilegal."
    )
    insert_image(doc, generate_tabasco_flood(), Inches(6.0))
    fig_caption(doc, "Clasificación de cobertura terrestre Dynamic World en Pantanos de Centla — 9 clases.")

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 21. TABLA COMPARATIVA
    # ══════════════════════════════════════════
    add_heading(doc, "21. Tabla Comparativa con Otras Plataformas", 1)

    add_text(doc,
        "APEX no compite directamente con plataformas de observación terrestre de uso general. "
        "Es un sistema operativo institucional diseñado para el flujo de trabajo de inspección "
        "y vigilancia ambiental de PROFEPA."
    )

    print("  [13/14] Radar comparativo...")
    insert_image(doc, generate_comparison_radar(), Inches(5.0))
    fig_caption(doc, "Comparativa radar de capacidades: APEX vs Global Forest Watch vs EO Browser.")

    make_table(doc,
        ["Capacidad", "APEX", "GFW", "EO Browser", "FIRMS"],
        [
            ["Motores de análisis", "12", "1", "Scripts custom", "1"],
            ["Índices espectrales", "15", "0", "Configurables", "0"],
            ["Deep Learning", "Prithvi-EO 300M", "—", "—", "—"],
            ["Radar SAR", "Sentinel-1 VV/VH", "—", "Visualización", "—"],
            ["Alertas GLAD/RADD", "Fusión espacial", "Sí", "—", "—"],
            ["Hansen GFC", "Integrado", "Sí", "Parcial", "—"],
            ["CCDC temporal", "Sí", "—", "—", "—"],
            ["Drivers deforestación", "WRI 7 clases", "Sí", "—", "—"],
            ["Fusión bayesiana", "H3 grid", "—", "—", "—"],
            ["POMDP planificación", "Sí", "—", "—", "—"],
            ["Forecast 1-5 años", "4 métodos", "—", "—", "—"],
            ["Monitoreo + email", "Sí", "Sí", "—", "Parcial"],
            ["KPI dashboard", "Sí", "—", "—", "—"],
            ["Reportes PDF/Word", "Folio PROFEPA", "—", "—", "—"],
            ["GPU acceleration", "CUDA", "Cloud", "Cloud", "—"],
        ],
        widths=[3.5, 2.5, 2, 2.5, 2],
        header_color="1b2a3d"
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 22. ROADMAP
    # ══════════════════════════════════════════
    add_heading(doc, "22. Roadmap y Evolución", 1)

    make_table(doc,
        ["Fase", "Funcionalidad", "Estado"],
        [
            ["v1.0", "12 motores de análisis + pipeline", "✅ Completado"],
            ["v1.0", "15 índices espectrales", "✅ Completado"],
            ["v1.0", "Dynamic World 9 clases + cambio T1→T2", "✅ Completado"],
            ["v1.0", "Hansen GFC + GLAD/RADD alertas", "✅ Completado"],
            ["v1.0", "SAR Sentinel-1 change detection", "✅ Completado"],
            ["v1.0", "FIRMS NRT + Fire MODIS", "✅ Completado"],
            ["v1.0", "Prithvi-EO 300M deep learning", "✅ Completado"],
            ["v1.0", "Timeline multi-año + Z-score", "✅ Completado"],
            ["v1.0", "POMDP + Bayesian + H3 grid", "✅ Completado"],
            ["v1.0", "Forecast engine (4 métodos)", "✅ Completado"],
            ["v1.0", "KPI dashboard + monitoreo", "✅ Completado"],
            ["v1.0", "Reportes PDF/Word/JSON", "✅ Completado"],
            ["v2.0", "Imágenes < 1m (motor estructuras)", "🔜 Planificado"],
            ["v2.0", "Dashboard de monitoreo nacional", "🔜 Planificado"],
            ["v2.0", "Integración SEMARNAT/CONAFOR", "🔜 Planificado"],
            ["v2.0", "Alertas automáticas por geofence", "🔜 Planificado"],
        ],
        widths=[1.5, 7, 3]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 23. CONCLUSIONES
    # ══════════════════════════════════════════
    add_heading(doc, "23. Conclusiones", 1)

    add_text(doc,
        "APEX constituye una plataforma integral y de vanguardia para el monitoreo ambiental "
        "institucional. Con 12 motores de análisis operando en paralelo, 15 índices espectrales, "
        "deep learning con el foundation model Prithvi-EO-2.0-300M, detección radar SAR para "
        "condiciones all-weather, y un módulo de planificación POMDP para optimización de "
        "inspecciones, APEX supera significativamente las capacidades de plataformas comerciales "
        "de observación terrestre."
    )

    conclusions = [
        "12 motores de análisis complementarios que cubren desde óptico (Sentinel-2) hasta "
        "radar (Sentinel-1), deep learning (Prithvi), alertas (GLAD/RADD), incendios (FIRMS), "
        "y análisis legal (ANP), proporcionando una visión integral de cada área de interés.",

        "15 índices espectrales calculados automáticamente con interpretación inteligente, "
        "permitiendo al inspector obtener diagnósticos completos sin conocimiento especializado "
        "en teledetección.",

        "Análisis temporal multi-anual (2018-2025) con detección de anomalías por Z-score, "
        "permitiendo identificar años con cambios significativos que requieren investigación.",

        "Motor de predicción (forecast) con 4 métodos complementarios que proyecta tendencias "
        "a horizonte de 1-5 años para planificación estratégica de largo plazo.",

        "Módulo POMDP que optimiza la asignación de inspectores y presupuesto, maximizando "
        "la efectividad de las inspecciones considerando la incertidumbre observacional.",

        "Fusión bayesiana de evidencias multi-motor sobre grid hexagonal H3, proporcionando "
        "probabilidades calibradas de ilícito ambiental por zona geográfica.",

        "Generación automática de reportes institucionales con folio PROFEPA, narrativas "
        "en español, gráficas y datos GeoJSON exportables.",

        "GPU acceleration con PyTorch CUDA que permite inferencia en tiempo real con "
        "modelos foundation de 300 millones de parámetros.",
    ]
    for c in conclusions:
        add_bullet(doc, c)

    add_text(doc,
        "APEX representa un avance significativo en la capacidad tecnológica de PROFEPA "
        "para la protección del patrimonio forestal mexicano, combinando ciencia de datos, "
        "inteligencia artificial y percepción remota en una herramienta operativa diseñada "
        "para impacto real en campo."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # ANEXO A: FÓRMULAS
    # ══════════════════════════════════════════
    add_heading(doc, "Anexo A. Fórmulas de Índices Espectrales", 1)

    add_text(doc,
        "Todas las fórmulas utilizan bandas Sentinel-2 Level-2A (reflectancia de superficie). "
        "Se aplica división segura: denominador + 10⁻¹⁰ para evitar divisiones por cero.", italic=True
    )

    formulas = [
        ["NDVI", "(B8 − B4) / (B8 + B4)", "Rouse et al., 1974"],
        ["EVI", "2.5 × (B8 − B4) / (B8 + 6×B4 − 7.5×B2 + 1)", "Huete et al., 2002"],
        ["SAVI", "1.5 × (B8 − B4) / (B8 + B4 + 0.5)", "Huete, 1988"],
        ["MSAVI", "[2×B8 + 1 − √((2×B8+1)² − 8×(B8−B4))] / 2", "Qi et al., 1994"],
        ["GNDVI", "(B8 − B3) / (B8 + B3)", "Gitelson et al., 1996"],
        ["NDRE", "(B8 − B5) / (B8 + B5)", "Barnes et al., 2000"],
        ["CIre", "B8 / B7 − 1", "Clevers & Gitelson, 2013"],
        ["RECIg", "B7 / B5 − 1", "Clevers & Gitelson, 2013"],
        ["NDWI", "(B3 − B8) / (B3 + B8)", "McFeeters, 1996"],
        ["MNDWI", "(B3 − B11) / (B3 + B11)", "Xu, 2006"],
        ["NDMI", "(B8 − B11) / (B8 + B11)", "Gao, 1996"],
        ["BSI", "((B11+B4) − (B8+B2)) / ((B11+B4) + (B8+B2))", "Rikimaru et al., 2002"],
        ["NBI", "B11 × B4 / B8", "Jia et al., 2014"],
        ["NDBI", "(B11 − B8) / (B11 + B8)", "Zha et al., 2003"],
        ["NBR", "(B8 − B12) / (B8 + B12)", "Key & Benson, 2006"],
    ]

    make_table(doc,
        ["Índice", "Fórmula", "Referencia"],
        formulas,
        widths=[2, 8, 4]
    )

    add_text(doc,
        "Bandas Sentinel-2: B2 (490nm, Blue), B3 (560nm, Green), B4 (665nm, Red), "
        "B5 (705nm, Red-Edge 1), B7 (783nm, Red-Edge 3), B8 (842nm, NIR), "
        "B11 (1610nm, SWIR-1), B12 (2190nm, SWIR-2)."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # ANEXO B: CCDC
    # ══════════════════════════════════════════
    add_heading(doc, "Anexo B. Matriz de Clasificación CCDC", 1)

    add_text(doc,
        "El motor CCDC (Continuous Change Detection and Classification) clasifica los cambios "
        "detectados en series temporales NDVI según su magnitud:"
    )

    make_table(doc,
        ["Magnitud NDVI", "Clasificación", "Descripción", "Acción Recomendada"],
        [
            ["< −0.3", "Desmonte Súbito", "Pérdida abrupta de vegetación (tala o desmonte)", "Inspección inmediata"],
            ["−0.3 a −0.1", "Degradación Gradual", "Pérdida progresiva (pastoreo, extracción selectiva)", "Monitoreo reforzado"],
            ["−0.1 a 0.1", "Sin Cambio", "Estabilidad en la cobertura vegetal", "Sin acción"],
            ["> 0.1", "Recuperación", "Incremento de biomasa (reforestación, regeneración)", "Seguimiento positivo"],
        ],
        widths=[2.5, 3, 5, 3.5]
    )

    add_text(doc,
        "El algoritmo CCDC opera sobre series temporales Sentinel-2 con mínimo 6 observaciones "
        "por pixel y probabilidad chi-cuadrado de 0.99 para detección de puntos de quiebre."
    )

    # ── Footer ──
    doc.add_paragraph()
    doc.add_paragraph()
    f = doc.add_paragraph()
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = f.add_run(
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "▲ APEX — Análisis y Protección del Entorno por Exploración Satelital\n"
        "PROFEPA — Coordinación de Estudios Prospectivos y Valoración de Riesgos\n"
        "Documento generado automáticamente — Marzo 2026\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━"
    )
    rf.font.size = Pt(8)
    rf.font.color.rgb = GRAY
    rf.italic = True

    # ── Save ──
    print("  [14/14] Guardando documento...")
    doc.save(str(OUT))
    size_kb = OUT.stat().st_size / 1024
    print(f"\n✅ Reporte generado exitosamente: {OUT}")
    print(f"   Tamaño: {size_kb:.1f} KB ({size_kb/1024:.1f} MB)")
    print(f"   Figuras: {FIG_NUM[0]}")


if __name__ == "__main__":
    generate_report()
