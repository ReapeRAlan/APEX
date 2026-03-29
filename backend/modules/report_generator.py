"""
Generador de Reportes PDF y Word — APEX (PROFEPA)
Adaptado del sistema MACOF para datos de timeline y análisis predictivo.

Genera reportes completos con:
  - Portada institucional
  - Resumen ejecutivo con indicadores clave
  - Alertas y anomalías (z-score > 2σ)
  - Datos tabulares por año
  - Gráficas de evolución temporal (barras + área apilada)
  - Mapas por año (polígonos de deforestación / expansión urbana)
  - Composición de vegetación por año (dona)
  - Interpretación automática y nivel de riesgo
  - Conclusiones y recomendaciones
  - Anexo metodológico
"""

import io
from datetime import datetime
from typing import Any, Dict

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.collections import PatchCollection
from matplotlib.patches import Polygon as MplPolygon
import numpy as np

# ── reportlab ──
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as RLImage,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

# ── python-docx ──
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

# ═══════════════════════════════════════════════════════════
#  Shared constants
# ═══════════════════════════════════════════════════════════

VEG_COLORS: Dict[str, str] = {
    "bosque_denso": "#166534",
    "bosque_ralo": "#22c55e",
    "pastizal": "#86efac",
    "manglar_inundado": "#7a87c6",
    "cultivos": "#e49635",
    "matorral": "#dfc35a",
    "urbano": "#6b21a8",
    "suelo": "#92400e",
    "agua": "#3b82f6",
    "quemado": "#7c2d12",
    "nieve": "#b39fe1",
}

VEG_LABELS: Dict[str, str] = {
    "bosque_denso": "Bosque denso",
    "bosque_ralo": "Bosque ralo",
    "pastizal": "Pastizal",
    "manglar_inundado": "Manglar / inundado",
    "cultivos": "Cultivos",
    "matorral": "Matorral",
    "urbano": "Urbano",
    "suelo": "Suelo desnudo",
    "agua": "Agua",
    "quemado": "Quemado",
    "nieve": "Nieve / hielo",
}


def _to_float(v: Any, default: float = 0.0) -> float:
    if v is None:
        return default
    try:
        return float(v)
    except (ValueError, TypeError):
        return default


def _fig_to_bytes(fig, dpi: int = 150) -> io.BytesIO:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight",
                facecolor=fig.get_facecolor(), edgecolor="none")
    plt.close(fig)
    buf.seek(0)
    return buf


# Optional heavy deps — graceful fallback
try:
    import geopandas as gpd
    from shapely.geometry import shape as shp_shape
    HAS_GEOPANDAS = True
except ImportError:
    HAS_GEOPANDAS = False

try:
    import contextily as cx
    HAS_CONTEXTILY = True
except ImportError:
    HAS_CONTEXTILY = False

try:
    from PIL import Image as PILImage
    HAS_PIL = True
except ImportError:
    HAS_PIL = False


# ═══════════════════════════════════════════════════════════
#  Cartographic helpers  (adapted from MACOF)
# ═══════════════════════════════════════════════════════════

def _overlay_satellite(ax, img_bytes: bytes, bounds: list,
                       alpha: float = 0.88, zorder: int = 2):
    """Overlay EE-generated satellite PNG onto a matplotlib axis."""
    if not HAS_PIL or not img_bytes:
        return
    try:
        pil_img = PILImage.open(io.BytesIO(img_bytes)).convert("RGBA")
        img_arr = np.array(pil_img)
        west, south, east, north = bounds
        ax.imshow(img_arr, extent=[west, east, south, north],
                  alpha=alpha, zorder=zorder, aspect="auto",
                  interpolation="bilinear")
    except Exception:
        pass


def _add_scale_bar(ax, coords: dict):
    """Publication-quality scale bar at lower-right."""
    import matplotlib.patheffects as pe
    try:
        xlim, ylim = ax.get_xlim(), ax.get_ylim()
        lat_center = (ylim[0] + ylim[1]) / 2
        km_per_deg = 111.32 * abs(np.cos(np.radians(lat_center)))
        map_width_km = (xlim[1] - xlim[0]) * km_per_deg
        nice = [0.01, 0.02, 0.05, 0.1, 0.2, 0.5, 1, 2, 5, 10, 20, 50, 100, 200, 500]
        bar_km = max(n for n in nice if n <= map_width_km * 0.35) if map_width_km > 0.02 else 0.1
        bar_deg = bar_km / km_per_deg

        x0 = xlim[1] - (xlim[1] - xlim[0]) * 0.05 - bar_deg
        y0 = ylim[0] + (ylim[1] - ylim[0]) * 0.06
        th = (ylim[1] - ylim[0]) * 0.012
        outline = [pe.withStroke(linewidth=4, foreground="white")]

        ax.plot([x0, x0 + bar_deg], [y0, y0], color="black", linewidth=4,
                solid_capstyle="butt", zorder=10, path_effects=outline)
        midx = x0 + bar_deg / 2
        ax.plot([x0, midx], [y0, y0], color="white", linewidth=2,
                solid_capstyle="butt", zorder=11)
        for x in [x0, midx, x0 + bar_deg]:
            ax.plot([x, x], [y0 - th, y0 + th], color="black", linewidth=1.5,
                    zorder=11, path_effects=outline)

        label = f"{bar_km:.0f} km" if bar_km >= 1 else f"{bar_km * 1000:.0f} m"
        ax.text(x0 + bar_deg / 2, y0 + th * 2.5, label, ha="center", va="bottom",
                fontsize=7, fontweight="bold", color="#1a1a1a", zorder=12,
                path_effects=[pe.withStroke(linewidth=3, foreground="white")])
    except Exception:
        pass


def _add_north_arrow(ax):
    """North arrow at upper-right."""
    import matplotlib.patheffects as pe
    try:
        xlim, ylim = ax.get_xlim(), ax.get_ylim()
        x = xlim[1] - (xlim[1] - xlim[0]) * 0.06
        y_top = ylim[1] - (ylim[1] - ylim[0]) * 0.05
        arrow_len = (ylim[1] - ylim[0]) * 0.09

        circle = plt.Circle((x, y_top - arrow_len * 0.3), arrow_len * 0.55,
                             facecolor="white", edgecolor="#94a3b8",
                             linewidth=0.8, alpha=0.85, zorder=9)
        ax.add_patch(circle)
        ax.annotate("", xy=(x, y_top), xytext=(x, y_top - arrow_len),
                     arrowprops=dict(arrowstyle="-|>", color="#1e3a5f",
                                     lw=2, mutation_scale=14),
                     zorder=10)
        ax.text(x, y_top + arrow_len * 0.12, "N", ha="center", va="bottom",
                fontsize=11, fontweight="bold", color="#1e3a5f", zorder=10,
                path_effects=[pe.withStroke(linewidth=3, foreground="white")])
    except Exception:
        pass


def _add_coord_grid(ax):
    """Coordinate grid with degree labels."""
    ax.grid(True, alpha=0.25, linestyle="--", linewidth=0.5, color="gray", zorder=1)
    ax.tick_params(labelsize=6.5, colors="#374151", direction="in", length=3, width=0.5)
    ax.xaxis.set_major_formatter(plt.FuncFormatter(
        lambda v, _: f'{abs(v):.4f}°{"W" if v < 0 else "E"}'))
    ax.yaxis.set_major_formatter(plt.FuncFormatter(
        lambda v, _: f'{abs(v):.4f}°{"S" if v < 0 else "N"}'))


def _geojson_to_gdf(geojson):
    """Parse any GeoJSON variant into a GeoDataFrame."""
    if not HAS_GEOPANDAS or not geojson:
        return None
    try:
        if isinstance(geojson, str):
            import json as _json
            geojson = _json.loads(geojson)
        gtype = geojson.get("type", "")
        if gtype == "FeatureCollection":
            gdf = gpd.GeoDataFrame.from_features(geojson["features"])
        elif gtype == "Feature":
            gdf = gpd.GeoDataFrame.from_features([geojson])
        elif gtype in ("Polygon", "MultiPolygon"):
            gdf = gpd.GeoDataFrame(geometry=[shp_shape(geojson)])
        else:
            return None
        return gdf.set_crs(epsg=4326)
    except Exception:
        return None


def _add_basemap(ax, gdf):
    """Add Esri satellite basemap tiles if contextily is available."""
    if not HAS_CONTEXTILY or gdf is None:
        return
    try:
        cx.add_basemap(ax, crs=gdf.crs, source=cx.providers.Esri.WorldImagery,
                       attribution_size=4, attribution="© Esri WorldImagery")
    except Exception:
        try:
            cx.add_basemap(ax, crs=gdf.crs, source=cx.providers.OpenStreetMap.Mapnik,
                           attribution_size=4)
        except Exception:
            pass


# ═══════════════════════════════════════════════════════════
#  Location overview map (new — like MACOF)
# ═══════════════════════════════════════════════════════════

def _build_location_map(aoi_geojson: dict, coords: dict,
                        area_ha: float = 0,
                        overview_thumb: dict = None) -> io.BytesIO:
    """Full-page location overview map with satellite basemap + polygon."""
    fig, ax = plt.subplots(figsize=(7, 5.5))
    fig.patch.set_facecolor("white")
    ax.set_facecolor("#eef2f6")

    gdf = _geojson_to_gdf(aoi_geojson)
    if gdf is not None and not gdf.empty:
        # Semi-transparent fill
        gdf.plot(ax=ax, facecolor="#00d4ff", alpha=0.15, edgecolor="none", zorder=3)
        # Strong boundary
        gdf.boundary.plot(ax=ax, color="#00d4ff", linewidth=2.5, zorder=5)

        # Set extent with padding
        minx, miny, maxx, maxy = gdf.total_bounds
        dx = max((maxx - minx) * 0.3, 0.005)
        dy = max((maxy - miny) * 0.3, 0.005)
        ax.set_xlim(minx - dx, maxx + dx)
        ax.set_ylim(miny - dy, maxy + dy)

        # Contextily basemap
        _add_basemap(ax, gdf)

        # Real satellite overlay if available
        if overview_thumb and overview_thumb.get("rgb"):
            _overlay_satellite(ax, overview_thumb["rgb"],
                               overview_thumb["bounds"], alpha=0.92, zorder=2)
            # Re-draw boundary on top
            gdf.boundary.plot(ax=ax, color="#00d4ff", linewidth=2.5, zorder=5)

    # Cartographic elements
    ax.set_title("Mapa de Ubicación del Área de Análisis",
                 fontsize=13, fontweight="bold", color="#1e3a5f", pad=14)
    _add_scale_bar(ax, coords)
    _add_north_arrow(ax)
    _add_coord_grid(ax)

    # Legend
    legend_items = [
        mpatches.Patch(facecolor="#00d4ff", alpha=0.3, edgecolor="#00d4ff",
                       linewidth=2, label=f"Área analizada: {area_ha:.2f} ha"),
    ]
    if coords.get("center_lat"):
        from matplotlib.lines import Line2D
        legend_items.append(Line2D(
            [0], [0], marker="o", color="w", markerfacecolor="#1e3a5f",
            markersize=6,
            label=f'{coords["center_lat"]:.4f}°N, {abs(coords.get("center_lon", 0)):.4f}°W'
        ))
    ax.legend(handles=legend_items, loc="lower left", fontsize=8,
              framealpha=0.9, edgecolor="#d1d5db", fancybox=True)

    # Footer
    fig.text(0.5, 0.01, "PROFEPA · APEX — Análisis Predictivo de Ecosistemas con IA",
             ha="center", va="bottom", fontsize=6, color="#6b7280", style="italic")
    fig.subplots_adjust(bottom=0.06)
    fig.tight_layout()
    return _fig_to_bytes(fig, dpi=180)


# ── Shared chart builders ──────────────────────────────────

def _build_change_bar_chart(timeline: dict) -> io.BytesIO:
    """Barras comparativas de deforestación vs expansión urbana por año."""
    years = sorted(timeline.keys())
    def_vals = [_to_float(timeline[y].get("deforestation", {}).get("stats", {}).get("area_ha")) for y in years]
    ue_vals = [_to_float(timeline[y].get("urban_expansion", {}).get("stats", {}).get("area_ha")) for y in years]

    fig, ax = plt.subplots(figsize=(6.5, 3))
    x = np.arange(len(years))
    w = 0.35
    bars1 = ax.bar(x - w / 2, def_vals, w, label="Deforestación (ha)", color="#ef4444", edgecolor="#b91c1c", linewidth=0.5)
    bars2 = ax.bar(x + w / 2, ue_vals, w, label="Exp. urbana (ha)", color="#f97316", edgecolor="#c2410c", linewidth=0.5)
    # Value labels on bars
    for bar in list(bars1) + list(bars2):
        h = bar.get_height()
        if h > 0:
            ax.text(bar.get_x() + bar.get_width() / 2, h + 0.05, f"{h:.1f}", ha="center", va="bottom", fontsize=7, color="#374151")
    ax.set_xticks(x)
    ax.set_xticklabels(years, fontsize=9, fontweight="bold")
    ax.set_ylabel("Hectáreas", fontsize=9)
    ax.set_title("Cambios detectados por año", fontsize=12, fontweight="bold", color="#1e3a5f")
    ax.legend(fontsize=8, loc="upper left")
    ax.grid(axis="y", alpha=0.3)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    return _fig_to_bytes(fig)


def _build_veg_stacked_area(timeline: dict) -> io.BytesIO:
    """Gráfica de área apilada: evolución de composición de vegetación."""
    years = sorted(timeline.keys())
    all_classes = set()
    for y in years:
        classes = timeline[y].get("vegetation", {}).get("stats", {}).get("classes", {})
        all_classes.update(k for k, v in classes.items() if v > 0)

    ordered = [c for c in VEG_COLORS if c in all_classes]
    data = {c: [] for c in ordered}
    for y in years:
        classes = timeline[y].get("vegetation", {}).get("stats", {}).get("classes", {})
        for c in ordered:
            data[c].append(_to_float(classes.get(c)))

    fig, ax = plt.subplots(figsize=(6.5, 3.5))
    bottom = np.zeros(len(years))
    for cls in ordered:
        vals = np.array(data[cls])
        ax.fill_between(range(len(years)), bottom, bottom + vals, label=VEG_LABELS.get(cls, cls),
                         color=VEG_COLORS.get(cls, "#6b7280"), alpha=0.85)
        bottom += vals

    ax.set_xticks(range(len(years)))
    ax.set_xticklabels(years, fontsize=9, fontweight="bold")
    ax.set_ylabel("Cobertura (%)", fontsize=9)
    ax.set_title("Evolución de la composición de vegetación", fontsize=12, fontweight="bold", color="#1e3a5f")
    ax.set_ylim(0, 100)
    ax.legend(fontsize=7, loc="center left", bbox_to_anchor=(1, 0.5), ncol=1)
    ax.grid(axis="y", alpha=0.2)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    return _fig_to_bytes(fig)


def _build_year_map(year_data: dict, year: str,
                    aoi_geojson: dict = None,
                    year_thumb: dict = None) -> io.BytesIO:
    """Cartographic map for a single year with satellite imagery overlay.

    If year_thumb is provided (from GEEThumbnailService), real Sentinel-2
    satellite imagery is overlaid behind the analysis polygons.
    Falls back to contextily basemap or plain colored polygons.
    """
    fig, ax = plt.subplots(figsize=(5, 4.5))
    fig.patch.set_facecolor("white")
    ax.set_facecolor("#eef2f6")

    has_data = False
    all_x, all_y = [], []

    # ── Try to render the AOI polygon as base ──
    gdf_aoi = _geojson_to_gdf(aoi_geojson) if aoi_geojson else None
    if gdf_aoi is not None and not gdf_aoi.empty:
        gdf_aoi.plot(ax=ax, facecolor="#00d4ff", alpha=0.08, edgecolor="none", zorder=1)
        gdf_aoi.boundary.plot(ax=ax, color="#00d4ff", linewidth=1.5,
                              linestyle="--", zorder=6, alpha=0.6)
        minx, miny, maxx, maxy = gdf_aoi.total_bounds
        all_x.extend([minx, maxx])
        all_y.extend([miny, maxy])

    def _plot_features(geojson: dict, color: str, label: str, alpha: float = 0.6):
        nonlocal has_data
        if not geojson or not geojson.get("features"):
            return
        patches = []
        for feat in geojson["features"]:
            geom = feat.get("geometry", {})
            if geom.get("type") == "Polygon":
                coords = geom["coordinates"][0]
                xs = [c[0] for c in coords]
                ys = [c[1] for c in coords]
                all_x.extend(xs)
                all_y.extend(ys)
                patches.append(MplPolygon(list(zip(xs, ys)), closed=True))
                has_data = True
            elif geom.get("type") == "MultiPolygon":
                for ring in geom["coordinates"]:
                    coords = ring[0]
                    xs = [c[0] for c in coords]
                    ys = [c[1] for c in coords]
                    all_x.extend(xs)
                    all_y.extend(ys)
                    patches.append(MplPolygon(list(zip(xs, ys)), closed=True))
                    has_data = True
        if patches:
            pc = PatchCollection(patches, alpha=alpha, facecolor=color, edgecolor=color, linewidth=0.8)
            ax.add_collection(pc)
            ax.plot([], [], color=color, label=label, linewidth=4, alpha=alpha)

    # Vegetation background
    veg_geo = year_data.get("vegetation", {}).get("geojson", {})
    if veg_geo and veg_geo.get("features"):
        for feat in veg_geo["features"]:
            geom = feat.get("geometry", {})
            cls = feat.get("properties", {}).get("class", "")
            c = VEG_COLORS.get(cls, "#d1d5db")
            if geom.get("type") == "Polygon":
                coords = geom["coordinates"][0]
                xs = [co[0] for co in coords]
                ys = [co[1] for co in coords]
                all_x.extend(xs)
                all_y.extend(ys)
                poly = MplPolygon(list(zip(xs, ys)), closed=True, facecolor=c, edgecolor="none", alpha=0.35)
                ax.add_patch(poly)

    _plot_features(year_data.get("deforestation", {}).get("geojson"), "#ef4444", "Deforestación", 0.7)
    _plot_features(year_data.get("urban_expansion", {}).get("geojson"), "#f97316", "Exp. urbana", 0.7)

    if not has_data and not all_x:
        ax.text(0.5, 0.5, "Sin datos geoespaciales", ha="center", va="center",
                fontsize=10, color="#9ca3af", transform=ax.transAxes)
    else:
        if all_x and all_y:
            dx = max(all_x) - min(all_x) or 0.01
            dy = max(all_y) - min(all_y) or 0.01
            pad = max(dx, dy) * 0.2
            ax.set_xlim(min(all_x) - pad, max(all_x) + pad)
            ax.set_ylim(min(all_y) - pad, max(all_y) + pad)

    # ── Satellite imagery overlay ──
    if year_thumb and year_thumb.get("rgb"):
        _overlay_satellite(ax, year_thumb["rgb"], year_thumb["bounds"],
                           alpha=0.85, zorder=2)
        # Re-draw analysis polygons on top after satellite
        _plot_features(year_data.get("deforestation", {}).get("geojson"), "#ef4444", "_skip", 0.5)
        _plot_features(year_data.get("urban_expansion", {}).get("geojson"), "#f97316", "_skip", 0.5)
        if gdf_aoi is not None and not gdf_aoi.empty:
            gdf_aoi.boundary.plot(ax=ax, color="#00d4ff", linewidth=1.5,
                                  linestyle="--", zorder=6, alpha=0.6)
    elif gdf_aoi is not None:
        _add_basemap(ax, gdf_aoi)

    baseline = year_data.get("baseline_year", int(year) - 1)
    ax.set_title(f"Análisis {baseline} → {year}", fontsize=11, fontweight="bold", color="#1e3a5f")

    # Cartographic elements
    coords = {}
    if all_y:
        coords["center_lat"] = (min(all_y) + max(all_y)) / 2
        coords["center_lon"] = (min(all_x) + max(all_x)) / 2
    _add_scale_bar(ax, coords)
    _add_north_arrow(ax)
    _add_coord_grid(ax)

    if has_data:
        # Filter duplicate labels
        handles, labels = ax.get_legend_handles_labels()
        unique = {}
        for h, lbl in zip(handles, labels):
            if lbl not in unique and not lbl.startswith("_"):
                unique[lbl] = h
        if unique:
            ax.legend(handles=list(unique.values()), labels=list(unique.keys()),
                      fontsize=7, loc="lower right", framealpha=0.9)
    ax.set_aspect("equal")
    fig.tight_layout()
    return _fig_to_bytes(fig, dpi=160)


def _build_veg_donut(classes: dict, year: str) -> io.BytesIO:
    """Gráfica de dona: composición de vegetación de un año."""
    filtered = {k: v for k, v in classes.items() if v > 0.5}
    if not filtered:
        filtered = {"Sin datos": 100}

    labels = [VEG_LABELS.get(k, k) for k in filtered]
    sizes = list(filtered.values())
    clrs = [VEG_COLORS.get(k, "#6b7280") for k in filtered]

    fig, ax = plt.subplots(figsize=(3.5, 3.5))
    wedges, texts, autotexts = ax.pie(
        sizes, labels=None, colors=clrs, autopct=lambda p: f"{p:.0f}%" if p >= 3 else "",
        startangle=90, pctdistance=0.78, wedgeprops=dict(width=0.38, edgecolor="white", linewidth=1.5),
    )
    for at in autotexts:
        at.set_fontsize(7)
        at.set_color("white")
        at.set_fontweight("bold")

    ax.legend(labels, fontsize=6.5, loc="center left", bbox_to_anchor=(1, 0.5))
    ax.set_title(f"Vegetación {year}", fontsize=10, fontweight="bold", color="#1e3a5f", pad=10)
    fig.tight_layout()
    return _fig_to_bytes(fig, dpi=130)


def _generate_interpretation(timeline: dict, cumulative: dict, anomalies: list) -> dict:
    """Auto-generate risk level, trend analysis, and recommendations."""
    total_def = _to_float(cumulative.get("total_deforestation_ha"))
    total_ue = _to_float(cumulative.get("total_urban_expansion_ha"))
    bosque_change = _to_float(cumulative.get("bosque_denso_change_pct"))
    n_anomalies = len(anomalies)
    n_years = _to_float(cumulative.get("years_analyzed", 1)) or 1

    # Risk score 0-100
    score = 0
    score += min(total_def * 3, 30)  # up to 30 pts for deforestation
    score += min(total_ue * 5, 20)   # up to 20 pts for urban expansion
    score += min(abs(bosque_change) * 2, 20)  # up to 20 pts for forest loss
    score += min(n_anomalies * 10, 20)  # up to 20 pts for anomalies
    score += min(total_def / n_years * 2, 10)  # rate factor
    score = min(int(score), 100)

    if score >= 70:
        level, level_color = "ALTO", "#ef4444"
    elif score >= 40:
        level, level_color = "MEDIO", "#f59e0b"
    else:
        level, level_color = "BAJO", "#22c55e"

    # Trend direction
    years = sorted(timeline.keys())
    def_series = [_to_float(timeline[y].get("deforestation", {}).get("stats", {}).get("area_ha")) for y in years]
    if len(def_series) >= 2:
        if def_series[-1] > def_series[-2] * 1.2:
            trend = "CRECIENTE — La deforestación muestra un incremento respecto al período anterior."
        elif def_series[-1] < def_series[-2] * 0.8:
            trend = "DECRECIENTE — Se observa una reducción respecto al período anterior."
        else:
            trend = "ESTABLE — Los niveles se mantienen relativamente constantes."
    else:
        trend = "INSUFICIENTE — Se requieren más años para determinar una tendencia."

    # Per-year detail summaries
    year_interpretations = {}
    for y in years:
        yd = timeline[y]
        d_ha = _to_float(yd.get("deforestation", {}).get("stats", {}).get("area_ha"))
        ue_ha = _to_float(yd.get("urban_expansion", {}).get("stats", {}).get("area_ha"))
        classes = yd.get("vegetation", {}).get("stats", {}).get("classes", {})
        bosque_pct = _to_float(classes.get("bosque_denso"))
        urbano_pct = _to_float(classes.get("urbano"))
        baseline = yd.get("baseline_year", int(y) - 1)

        lines = []
        if d_ha > 0:
            n_feat = yd.get("deforestation", {}).get("stats", {}).get("n_features", 0)
            conf = _to_float(yd.get("deforestation", {}).get("stats", {}).get("confidence"))
            lines.append(f"Se detectaron {d_ha:.1f} ha de deforestación en {n_feat} zona(s) "
                         f"(confianza promedio: {conf*100:.0f}%) respecto al año base {baseline}.")
        if ue_ha > 0:
            n_feat = yd.get("urban_expansion", {}).get("stats", {}).get("n_features", 0)
            lines.append(f"Expansión urbana de {ue_ha:.1f} ha detectada en {n_feat} zona(s).")
            # Check for alerts
            feats = yd.get("urban_expansion", {}).get("geojson", {}).get("features", [])
            alertas = [f.get("properties", {}).get("alerta", "") for f in feats if f.get("properties", {}).get("alerta")]
            if alertas:
                lines.append(f"⚠ Alertas: {'; '.join(set(alertas))}")
        if bosque_pct > 0:
            lines.append(f"Cobertura de bosque denso: {bosque_pct:.1f}%, urbano: {urbano_pct:.1f}%.")
        if not lines:
            lines.append("Sin cambios significativos detectados en este período.")

        year_interpretations[y] = " ".join(lines)

    # Recommendations
    recommendations = []
    if total_def > 5:
        recommendations.append("Verificar en campo las zonas con mayor pérdida de cobertura forestal.")
    if total_ue > 3:
        recommendations.append("Revisar permisos de construcción en las zonas de expansión urbana detectadas.")
    if n_anomalies > 0:
        recommendations.append("Investigar las anomalías estadísticas detectadas — podrían indicar eventos extraordinarios.")
    if bosque_change < -5:
        recommendations.append("La tendencia de pérdida de bosque denso requiere atención prioritaria y posible intervención.")
    if not recommendations:
        recommendations.append("Continuar el monitoreo periódico. Los indicadores se encuentran dentro de parámetros normales.")

    return {
        "score": score,
        "level": level,
        "level_color": level_color,
        "trend": trend,
        "year_interpretations": year_interpretations,
        "recommendations": recommendations,
    }


# ═══════════════════════════════════════════════════════════
#  PDF Report Generator
# ═══════════════════════════════════════════════════════════

class APEXPDFReportGenerator:
    """Genera reportes PDF institucionales completos para APEX / PROFEPA."""

    # Colores institucionales
    COLOR_PRIMARY = colors.HexColor("#00d4ff")
    COLOR_SECONDARY = colors.HexColor("#1e3a5f")
    COLOR_GREEN = colors.HexColor("#2ea043")
    COLOR_ORANGE = colors.HexColor("#f0883e")
    COLOR_RED = colors.HexColor("#ef4444")
    COLOR_TEXT = colors.HexColor("#1f2937")

    def __init__(self):
        self.styles = self._create_styles()

    def _create_styles(self) -> Dict[str, Any]:
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name="CoverTitle", parent=styles["Heading1"],
            fontSize=26, textColor=colors.white,
            alignment=TA_CENTER, fontName="Helvetica-Bold", spaceAfter=20,
        ))
        styles.add(ParagraphStyle(
            name="CoverSub", parent=styles["Normal"],
            fontSize=12, textColor=colors.HexColor("#b0c4de"),
            alignment=TA_CENTER, fontName="Helvetica",
        ))
        styles.add(ParagraphStyle(
            name="SectionHead", parent=styles["Heading2"],
            fontSize=15, textColor=self.COLOR_SECONDARY,
            spaceAfter=10, spaceBefore=16, fontName="Helvetica-Bold",
        ))
        styles.add(ParagraphStyle(
            name="SubHead", parent=styles["Heading3"],
            fontSize=12, textColor=colors.HexColor("#374151"),
            spaceAfter=6, spaceBefore=10, fontName="Helvetica-Bold",
        ))
        styles.add(ParagraphStyle(
            name="Body", parent=styles["BodyText"],
            fontSize=10, textColor=self.COLOR_TEXT,
            alignment=TA_JUSTIFY, spaceAfter=8, leading=13,
        ))
        styles.add(ParagraphStyle(
            name="BodySmall", parent=styles["BodyText"],
            fontSize=9, textColor=colors.HexColor("#4b5563"),
            alignment=TA_JUSTIFY, spaceAfter=6, leading=12,
        ))
        styles.add(ParagraphStyle(
            name="Small", parent=styles["Normal"],
            fontSize=8, textColor=colors.HexColor("#6b7280"),
            alignment=TA_CENTER,
        ))
        styles.add(ParagraphStyle(
            name="RiskHigh", parent=styles["Normal"],
            fontSize=14, textColor=colors.white,
            alignment=TA_CENTER, fontName="Helvetica-Bold",
        ))
        return styles

    def _header_footer(self, canvas_obj, doc):
        canvas_obj.saveState()
        canvas_obj.setFillColor(self.COLOR_SECONDARY)
        canvas_obj.rect(0, letter[1] - 0.7 * inch, letter[0], 0.7 * inch, fill=True, stroke=False)
        canvas_obj.setFillColor(colors.white)
        canvas_obj.setFont("Helvetica-Bold", 12)
        canvas_obj.drawString(0.75 * inch, letter[1] - 0.45 * inch, "PROFEPA — APEX")
        canvas_obj.setFont("Helvetica", 8)
        canvas_obj.drawRightString(
            letter[0] - 0.75 * inch, letter[1] - 0.45 * inch,
            f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        )
        canvas_obj.setStrokeColor(self.COLOR_PRIMARY)
        canvas_obj.setLineWidth(0.5)
        canvas_obj.line(0.75 * inch, 0.55 * inch, letter[0] - 0.75 * inch, 0.55 * inch)
        canvas_obj.setFont("Helvetica", 7)
        canvas_obj.setFillColor(self.COLOR_TEXT)
        canvas_obj.drawCentredString(letter[0] / 2, 0.38 * inch, f"Página {canvas_obj.getPageNumber()}")
        canvas_obj.drawString(0.75 * inch, 0.38 * inch, "Análisis Predictivo de Ecosistemas con IA")
        canvas_obj.drawRightString(letter[0] - 0.75 * inch, 0.38 * inch, "Confidencial")
        canvas_obj.restoreState()

    # ── 0. Cover ──
    def _build_cover(self, folio: str, period: str, interp: dict) -> list:
        story: list = []
        story.append(Spacer(1, 1.5 * inch))
        cover_table = Table(
            [[Paragraph("REPORTE DE ANÁLISIS<br/>PREDICTIVO ECOSISTÉMICO", self.styles["CoverTitle"])]],
            colWidths=[6.5 * inch], rowHeights=[2.5 * inch],
        )
        cover_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), self.COLOR_SECONDARY),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("LEFTPADDING", (0, 0), (-1, -1), 20),
            ("RIGHTPADDING", (0, 0), (-1, -1), 20),
        ]))
        story.append(cover_table)
        story.append(Spacer(1, 0.4 * inch))

        # Risk badge
        risk_color = colors.HexColor(interp["level_color"])
        risk_table = Table(
            [[Paragraph(f"NIVEL DE RIESGO: {interp['level']}  ({interp['score']}/100)", self.styles["RiskHigh"])]],
            colWidths=[4 * inch], rowHeights=[0.5 * inch],
        )
        risk_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), risk_color),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("ROUNDEDCORNERS", [6, 6, 6, 6]),
        ]))
        # Center the risk badge
        outer = Table([[risk_table]], colWidths=[6.5 * inch])
        outer.setStyle(TableStyle([("ALIGN", (0, 0), (-1, -1), "CENTER")]))
        story.append(outer)
        story.append(Spacer(1, 0.4 * inch))

        meta_text = (
            f"<b>Folio:</b> {folio}<br/>"
            f"<b>Período:</b> {period}<br/>"
            f"<b>Fecha de generación:</b> {datetime.now().strftime('%Y-%m-%d %H:%M UTC')}<br/>"
            f"<b>Sistema:</b> APEX — Análisis Predictivo de Ecosistemas con IA<br/>"
            f"<b>Institución:</b> PROFEPA"
        )
        story.append(Paragraph(meta_text, self.styles["Body"]))
        story.append(PageBreak())
        return story

    # ── 1. Executive summary ──
    def _build_summary(self, cumulative: dict, n_anomalies: int, interp: dict) -> list:
        story: list = []
        story.append(Paragraph("1. RESUMEN EJECUTIVO", self.styles["SectionHead"]))

        data = [
            ["Indicador", "Valor", "Interpretación"],
            ["Deforestación acumulada",
             f"{cumulative.get('total_deforestation_ha', 0)} ha",
             "Alta" if _to_float(cumulative.get("total_deforestation_ha")) > 10 else "Normal"],
            ["Expansión urbana total",
             f"{cumulative.get('total_urban_expansion_ha', 0)} ha",
             "Alta" if _to_float(cumulative.get("total_urban_expansion_ha")) > 5 else "Normal"],
            ["Cambio bosque denso",
             f"{cumulative.get('bosque_denso_change_pct', 0)} pp",
             "Pérdida" if _to_float(cumulative.get("bosque_denso_change_pct")) < -1 else "Estable"],
            ["Cambio urbano",
             f"{cumulative.get('urbano_change_pct', 0)} pp",
             "Creciendo" if _to_float(cumulative.get("urbano_change_pct")) > 1 else "Estable"],
            ["Años analizados", str(cumulative.get("years_analyzed", "-")), ""],
            ["Anomalías detectadas", str(n_anomalies),
             "⚠ Requiere atención" if n_anomalies > 0 else "Sin alertas"],
        ]
        t = Table(data, colWidths=[2.5 * inch, 1.8 * inch, 2.2 * inch])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), self.COLOR_SECONDARY),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d1d5db")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f4f8")]),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        story.append(t)
        story.append(Spacer(1, 0.15 * inch))

        # Trend text
        story.append(Paragraph(f"<b>Tendencia general:</b> {interp['trend']}", self.styles["Body"]))
        return story

    # ── 1b. Legal context (ANP / WDPA) ──
    def _build_legal_context(self, summary: dict) -> list:
        """Sección de contexto legal: intersección con ANPs (WDPA)."""
        lc = summary.get("legal_context")
        if not lc:
            return []

        story: list = []
        story.append(Spacer(1, 0.2 * inch))
        story.append(Paragraph("CONTEXTO LEGAL — ÁREAS NATURALES PROTEGIDAS", self.styles["SectionHead"]))

        anp_name = lc.get("anp_name", "Desconocida")
        category = lc.get("category", "N/D")
        overlap_area = _to_float(lc.get("overlap_area_ha"))
        overlap_pct = _to_float(lc.get("overlap_pct"))
        risk_assessment = lc.get("risk_assessment", "No disponible")
        intersects = lc.get("intersects_anp", False)

        data = [
            ["Parámetro", "Valor"],
            ["Nombre del ANP", anp_name],
            ["Categoría", category],
            ["Área de traslape", f"{overlap_area:.2f} ha"],
            ["Porcentaje de traslape", f"{overlap_pct:.1f}%"],
            ["Evaluación de riesgo", risk_assessment],
        ]
        t = Table(data, colWidths=[2.5 * inch, 4 * inch])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), self.COLOR_SECONDARY),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d1d5db")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f4f8")]),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        story.append(t)

        if intersects:
            story.append(Spacer(1, 0.1 * inch))
            warning_text = (
                "<b>ADVERTENCIA:</b> El área de interés se traslapa con un Área Natural Protegida. "
                "La deforestación dentro de una ANP puede constituir una infracción conforme al "
                "<b>Artículo 47 de la LGEEPA</b> (Ley General del Equilibrio Ecológico y la "
                "Protección al Ambiente). Se recomienda verificación en campo y, en su caso, "
                "iniciar procedimiento administrativo."
            )
            warning_style = ParagraphStyle(
                "WarningBox", parent=self.styles["Body"],
                textColor=colors.white, fontSize=9, leading=12,
            )
            warning_para = Paragraph(warning_text, warning_style)
            warning_table = Table([[warning_para]], colWidths=[6.5 * inch])
            warning_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), self.COLOR_RED),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ("ROUNDEDCORNERS", [4, 4, 4, 4]),
            ]))
            story.append(warning_table)

        return story

    # ── 1c. Hansen historical reference ──
    def _build_hansen_section(self, summary: dict) -> list:
        """Sección Hansen: referencia histórica de pérdida forestal (2001-2024)."""
        hansen = summary.get("hansen")
        if not hansen:
            return []

        story: list = []
        story.append(Spacer(1, 0.2 * inch))
        story.append(Paragraph("HANSEN — REFERENCIA HISTÓRICA (2001-2024)", self.styles["SectionHead"]))

        total_loss = _to_float(hansen.get("total_loss_ha"))
        n_features = hansen.get("n_features", 0)
        avg_treecover = _to_float(hansen.get("avg_treecover_pct"))

        intro = (
            f"Según el dataset <b>Hansen Global Forest Change v1.12</b> (UMD), se registra una "
            f"pérdida forestal acumulada de <b>{total_loss:.2f} ha</b> distribuida en "
            f"<b>{n_features}</b> polígono(s), con una cobertura arbórea promedio inicial de "
            f"<b>{avg_treecover:.1f}%</b>."
        )
        story.append(Paragraph(intro, self.styles["Body"]))

        # Loss by year table
        loss_by_year = hansen.get("loss_by_year", {})
        if loss_by_year:
            story.append(Spacer(1, 0.1 * inch))
            story.append(Paragraph("Desglose de pérdida por año:", self.styles["SubHead"]))
            rows = [["Año", "Pérdida (ha)"]]
            for yr in sorted(loss_by_year.keys()):
                rows.append([str(yr), f"{_to_float(loss_by_year[yr]):.2f}"])
            t = Table(rows, colWidths=[2 * inch, 2 * inch])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), self.COLOR_SECONDARY),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d1d5db")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f4f8")]),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (1, 1), (1, -1), "CENTER"),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(t)

        story.append(Spacer(1, 0.05 * inch))
        story.append(Paragraph(
            "Fuente: UMD / Hansen — Global Forest Change v1.12 (resolución 30 m)",
            self.styles["Small"],
        ))
        return story

    # ── 1d. Drivers de deforestación ──
    def _build_drivers_section(self, summary: dict) -> list:
        """Sección de drivers de deforestación (WRI / Google DeepMind)."""
        drivers = summary.get("drivers")
        if not drivers:
            return []

        story: list = []
        story.append(Spacer(1, 0.2 * inch))
        story.append(Paragraph("CLASIFICACIÓN DE DRIVERS DE DEFORESTACIÓN", self.styles["SectionHead"]))

        story.append(Paragraph(
            "Distribución de las causas de deforestación identificadas en el área de interés, "
            "según la clasificación de <b>WRI / Google DeepMind</b>:",
            self.styles["Body"],
        ))

        # Build distribution table
        distribution = drivers.get("distribution", drivers.get("classes",
                       drivers.get("drivers", {})))
        if isinstance(distribution, list):
            # Handle list-of-dicts format: [{"driver": "...", "percentage": ...}, ...]
            rows = [["Driver", "Porcentaje"]]
            for item in distribution:
                name = item.get("driver", item.get("name", "Desconocido"))
                pct = _to_float(item.get("percentage", item.get("pct")))
                rows.append([name, f"{pct:.1f}%"])
        elif isinstance(distribution, dict):
            rows = [["Driver", "Porcentaje"]]
            for driver_name in sorted(distribution.keys()):
                pct = _to_float(distribution[driver_name])
                rows.append([driver_name, f"{pct:.1f}%"])
        else:
            return story  # No parseable distribution

        if len(rows) > 1:
            t = Table(rows, colWidths=[3.5 * inch, 2 * inch])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), self.COLOR_SECONDARY),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d1d5db")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f4f8")]),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (1, 1), (1, -1), "CENTER"),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(t)

        story.append(Spacer(1, 0.05 * inch))
        story.append(Paragraph(
            "Fuente: WRI / Google DeepMind — Drivers de Deforestación (resolución 1 km)",
            self.styles["Small"],
        ))
        return story

    # ── 1d. Fire / Burned Area section (PDF) ──
    def _build_fire_section(self, summary: dict) -> list:
        """Sección de incendios / áreas quemadas — PDF."""
        fire = summary.get("fire")
        if not fire:
            return []
        story: list = []
        story.append(Spacer(1, 0.2 * inch))
        story.append(Paragraph("Incendios y Áreas Quemadas (MODIS)", self.styles["SectionHead"]))

        total_burned = _to_float(fire.get("total_burned_ha"))
        fire_count = fire.get("fire_count", 0)
        fire_defor_pct = _to_float(fire.get("fire_related_deforestation_pct"))

        story.append(Paragraph(
            f"Se detectaron <b>{fire_count}</b> zona(s) quemada(s) con un área total de "
            f"<b>{total_burned:.2f} ha</b> según el sensor MODIS MCD64A1 (resolución 500 m). "
            f"Se estima que el <b>{fire_defor_pct:.1f}%</b> de la deforestación detectada "
            f"está asociada espacialmente con áreas quemadas.",
            self.styles["Body"],
        ))
        story.append(Spacer(1, 0.05 * inch))
        story.append(Paragraph(
            "Fuente: MODIS MCD64A1 — Burned Area (resolución 500 m, composición mensual)",
            self.styles["Small"],
        ))
        return story

    # ── 1e. Deforestation Alerts GLAD/RADD section (PDF) ──
    def _build_deforestation_alerts_section(self, summary: dict) -> list:
        """Sección de alertas de deforestación GLAD/RADD — PDF."""
        alerts = summary.get("alerts")
        if not alerts:
            return []
        story: list = []
        story.append(Spacer(1, 0.2 * inch))
        story.append(Paragraph("Alertas de Deforestación (GLAD / RADD)", self.styles["SectionHead"]))

        total_alerts = alerts.get("total_alerts", 0)
        glad_count = alerts.get("glad_count", 0)
        radd_count = alerts.get("radd_count", 0)
        confirmed = alerts.get("confirmed_count", 0)
        total_area = _to_float(alerts.get("total_area_ha"))

        story.append(Paragraph(
            f"Se identificaron <b>{total_alerts}</b> alertas de deforestación activas: "
            f"<b>{glad_count}</b> GLAD (Sentinel-2 óptico) y <b>{radd_count}</b> RADD "
            f"(Sentinel-1 SAR, penetra nubes), cubriendo un área total de <b>{total_area:.2f} ha</b>. "
            f"De estas, <b>{confirmed}</b> tienen confianza alta/confirmada.",
            self.styles["Body"],
        ))

        if radd_count > 0 and glad_count == 0:
            story.append(Paragraph(
                "<b>Nota:</b> Solo se detectaron alertas SAR (RADD), lo que sugiere "
                "posible cobertura de nubes persistente que impide la detección óptica.",
                self.styles["BodySmall"],
            ))

        story.append(Spacer(1, 0.05 * inch))
        story.append(Paragraph(
            "Fuente: GLAD-S2 (UMD, 10 m óptico) + RADD (WUR, 10 m SAR)",
            self.styles["Small"],
        ))
        return story

    # ── 1f. SAR Change Detection section (PDF) ──
    def _build_sar_section(self, summary: dict) -> list:
        """Sección de detección de cambios SAR — PDF."""
        sar = summary.get("sar")
        if not sar:
            return []
        story: list = []
        story.append(Spacer(1, 0.2 * inch))
        story.append(Paragraph("Detección de Cambios SAR (Sentinel-1)", self.styles["SectionHead"]))

        n_changes = sar.get("sar_change_count", 0)
        total_area = _to_float(sar.get("total_change_ha", sar.get("total_area_ha")))
        high_conf = sar.get("high_confidence_count", 0)

        story.append(Paragraph(
            f"Mediante análisis bi-temporal de log-ratio sobre composites radar "
            f"Sentinel-1 (banda C, VV+VH), se detectaron <b>{n_changes}</b> "
            f"zonas de cambio con un área total de <b>{total_area:.2f} ha</b>. "
            f"De estas, <b>{high_conf}</b> presentan alta confianza (caída &gt;5 dB). "
            f"La detección SAR es complementaria a la óptica y funciona al 100% bajo "
            f"cobertura de nubes.",
            self.styles["Body"],
        ))
        story.append(Spacer(1, 0.05 * inch))
        story.append(Paragraph(
            "Fuente: Sentinel-1 GRD IW — Log-ratio change detection (resolución 10 m)",
            self.styles["Small"],
        ))
        return story

    # ── 1g. Cross-validation section (PDF) ──
    def _build_crossval_section(self, summary: dict) -> list:
        """Sección de validación cruzada DW vs MapBiomas — PDF."""
        cv = summary.get("crossval")
        if not cv:
            return []
        story: list = []
        story.append(Spacer(1, 0.2 * inch))
        story.append(Paragraph("Validación Cruzada (Dynamic World vs MapBiomas)", self.styles["SectionHead"]))

        agreement = _to_float(cv.get("agreement_pct"))
        disagreement = cv.get("disagreement_zones", 0)
        total = cv.get("total_compared", 0)

        story.append(Paragraph(
            f"Se compararon <b>{total}</b> detecciones de deforestación (Dynamic World) contra "
            f"la clasificación LULC de MapBiomas México. La tasa de acuerdo es del "
            f"<b>{agreement:.1f}%</b>, con <b>{disagreement}</b> zona(s) de desacuerdo "
            f"identificadas como posibles falsos positivos.",
            self.styles["Body"],
        ))

        if agreement >= 80:
            story.append(Paragraph(
                "La alta tasa de acuerdo refuerza la confiabilidad de las detecciones.",
                self.styles["BodySmall"],
            ))
        elif agreement < 60:
            story.append(Paragraph(
                "<b>Nota:</b> La tasa de acuerdo moderada sugiere prudencia en la "
                "interpretación. Se recomienda verificación en campo.",
                self.styles["BodySmall"],
            ))

        story.append(Spacer(1, 0.05 * inch))
        story.append(Paragraph(
            "Fuente: MapBiomas México v1.0 (resolución 30 m)",
            self.styles["Small"],
        ))
        return story

    # ── 1g. FIRMS Hotspots section (PDF) ──
    def _build_firms_section(self, summary: dict) -> list:
        """Sección de puntos de calor FIRMS — PDF."""
        # Aggregate FIRMS from timeline per-year data
        timeline = summary.get("timeline", {})
        cumulative = summary.get("cumulative", {})
        total_hotspots = cumulative.get("total_firms_hotspots", 0)
        total_frp = _to_float(cumulative.get("total_frp_mw", 0))

        if total_hotspots == 0:
            return []

        story: list = []
        story.append(Spacer(1, 0.2 * inch))
        story.append(Paragraph("Puntos de Calor FIRMS (VIIRS / MODIS NRT)", self.styles["SectionHead"]))

        story.append(Paragraph(
            f"Se identificaron <b>{total_hotspots}</b> detecciones de puntos de calor activos "
            f"en el área de estudio durante el período analizado, con una potencia radiativa "
            f"total de <b>{total_frp:,.1f} MW</b> (Fire Radiative Power). "
            f"Los datos provienen del sistema NASA FIRMS con sensores VIIRS (375 m) y MODIS (1 km).",
            self.styles["Body"],
        ))

        # Per-year table
        table_data = [["Año", "Hotspots", "Alta conf.", "Clusters", "FRP total (MW)", "FRP máx (MW)", "Satélites"]]
        for yr in sorted(timeline.keys()):
            yr_firms = timeline[yr].get("firms_hotspots", {}).get("stats", {})
            hc = yr_firms.get("hotspot_count", 0)
            if hc == 0:
                continue
            hi = yr_firms.get("high_confidence_count", 0)
            cl = yr_firms.get("cluster_count", 0)
            frp_t = yr_firms.get("total_frp_mw", 0)
            frp_m = yr_firms.get("max_frp_mw", 0)
            sats = ", ".join(yr_firms.get("satellites", []))
            table_data.append([str(yr), str(hc), str(hi), str(cl),
                               f"{frp_t:,.1f}", f"{frp_m:,.1f}", sats])

        if len(table_data) > 1:
            story.append(Spacer(1, 0.1 * inch))
            col_w = [0.6 * inch, 0.7 * inch, 0.7 * inch, 0.7 * inch, 1.0 * inch, 1.0 * inch, 1.3 * inch]
            t = Table(table_data, colWidths=col_w, repeatRows=1)
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a1a2e")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, 0), 7),
                ("FONTSIZE", (0, 1), (-1, -1), 6.5),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 4),
                ("GRID", (0, 0), (-1, -1), 0.3, colors.grey),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#f8f8f8"), colors.white]),
                ("ALIGN", (1, 0), (-1, -1), "CENTER"),
            ]))
            story.append(t)

        story.append(Spacer(1, 0.1 * inch))
        story.append(Paragraph(
            "<b>FRP (Fire Radiative Power)</b>: Potencia radiativa del fuego medida en megawatts. "
            "Valores altos indican mayor intensidad de quema y potencial de daño.",
            self.styles["BodySmall"],
        ))
        story.append(Paragraph(
            "Fuente: NASA FIRMS — VIIRS (S-NPP, NOAA-20) 375 m / MODIS (Terra, Aqua) 1 km, datos NRT + archivo",
            self.styles["Small"],
        ))
        return story

    # ── 2. Alerts ──
    def _build_alerts(self, anomalies: list) -> list:
        story: list = []
        story.append(Spacer(1, 0.2 * inch))
        story.append(Paragraph("2. ALERTAS Y ANOMALÍAS", self.styles["SectionHead"]))
        if not anomalies:
            story.append(Paragraph(
                "No se detectaron anomalías estadísticas (z-score &gt; 2σ) en la serie temporal analizada. "
                "Todos los indicadores se encuentran dentro de la variabilidad histórica esperada.",
                self.styles["Body"],
            ))
            return story

        story.append(Paragraph(
            f"Se detectaron <b>{len(anomalies)}</b> anomalía(s) estadística(s) en la serie temporal. "
            "Una anomalía indica que el valor observado excede 2 desviaciones estándar respecto a la media histórica.",
            self.styles["Body"],
        ))

        header = ["Año", "Motor", "Valor (ha)", "Media hist.", "z-score", "Severidad"]
        rows = [header]
        for a in anomalies:
            rows.append([
                str(a.get("year", "")),
                a.get("engine", "").replace("_", " ").title(),
                f"{_to_float(a.get('value', a.get('area_ha'))):.1f}",
                f"{_to_float(a.get('mean_ha')):.1f}",
                f"{_to_float(a.get('z_score')):.2f}",
                a.get("severity", "").upper(),
            ])
        t = Table(rows, colWidths=[0.8 * inch, 1.3 * inch, 1.1 * inch, 1.1 * inch, 0.9 * inch, 1.3 * inch])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), self.COLOR_RED),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d1d5db")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#fef2f2")]),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        story.append(t)

        # Interpretation per anomaly
        for a in anomalies:
            msg = a.get("message", "")
            if msg:
                story.append(Paragraph(f"• {msg}", self.styles["BodySmall"]))
        return story

    # ── 3. Annual data table ──
    def _build_annual_data(self, timeline: dict) -> list:
        story: list = []
        story.append(Spacer(1, 0.2 * inch))
        story.append(Paragraph("3. DATOS TABULARES POR AÑO", self.styles["SectionHead"]))

        header = ["Año", "Base", "Def. (ha)", "Exp. Urb. (ha)", "Bosque (%)", "Urbano (%)", "Confianza"]
        rows = [header]
        for year_str in sorted(timeline.keys()):
            yd = timeline[year_str]
            def_stats = yd.get("deforestation", {}).get("stats", {})
            ue_stats = yd.get("urban_expansion", {}).get("stats", {})
            classes = yd.get("vegetation", {}).get("stats", {}).get("classes", {})
            conf = _to_float(def_stats.get("confidence"))
            rows.append([
                year_str,
                str(yd.get("baseline_year", "?")),
                f"{_to_float(def_stats.get('area_ha')):.1f}",
                f"{_to_float(ue_stats.get('area_ha')):.1f}",
                f"{_to_float(classes.get('bosque_denso')):.1f}",
                f"{_to_float(classes.get('urbano')):.1f}",
                f"{conf*100:.0f}%" if conf else "—",
            ])
        t = Table(rows, colWidths=[0.7*inch, 0.7*inch, 1*inch, 1.1*inch, 1*inch, 1*inch, 1*inch])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), self.COLOR_SECONDARY),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d1d5db")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f4f8")]),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("ALIGN", (2, 1), (-1, -1), "CENTER"),
        ]))
        story.append(t)
        return story

    # ── 4. Charts ──
    def _build_charts(self, timeline: dict) -> list:
        story: list = []
        story.append(PageBreak())
        story.append(Paragraph("4. GRÁFICAS DE EVOLUCIÓN TEMPORAL", self.styles["SectionHead"]))

        # Bar chart
        story.append(Paragraph("4.1 Cambios detectados por año", self.styles["SubHead"]))
        buf = _build_change_bar_chart(timeline)
        story.append(RLImage(buf, width=6.5 * inch, height=3 * inch))
        story.append(Spacer(1, 0.2 * inch))

        # Stacked area
        story.append(Paragraph("4.2 Evolución de la composición de vegetación", self.styles["SubHead"]))
        buf2 = _build_veg_stacked_area(timeline)
        story.append(RLImage(buf2, width=6.5 * inch, height=3.5 * inch))
        return story

    # ── 5. Year-by-year detail pages ──
    def _build_year_pages(self, timeline: dict, interp: dict,
                          aoi_geojson: dict = None,
                          thumbnails: dict = None) -> list:
        thumbnails = thumbnails or {}
        story: list = []
        story.append(PageBreak())
        story.append(Paragraph("5. ANÁLISIS DETALLADO POR AÑO", self.styles["SectionHead"]))

        for year_str in sorted(timeline.keys()):
            yd = timeline[year_str]
            baseline = yd.get("baseline_year", int(year_str) - 1)

            story.append(Spacer(1, 0.15 * inch))
            story.append(Paragraph(
                f"5.{year_str[-2:]} Año {year_str} (base: {baseline})",
                self.styles["SubHead"],
            ))

            # Interpretation text
            txt = interp["year_interpretations"].get(year_str, "")
            if txt:
                story.append(Paragraph(txt, self.styles["BodySmall"]))

            # Side-by-side: Cartographic Map + donut
            year_thumb = thumbnails.get(year_str)
            map_buf = _build_year_map(yd, year_str,
                                      aoi_geojson=aoi_geojson,
                                      year_thumb=year_thumb)
            classes = yd.get("vegetation", {}).get("stats", {}).get("classes", {})
            donut_buf = _build_veg_donut(classes, year_str)

            img_table = Table(
                [[RLImage(map_buf, width=3.3 * inch, height=3 * inch),
                  RLImage(donut_buf, width=3 * inch, height=3 * inch)]],
                colWidths=[3.35 * inch, 3.15 * inch],
            )
            img_table.setStyle(TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ]))
            story.append(img_table)

            # ── NDVI satellite thumbnail (if available) ──
            if year_thumb and year_thumb.get("ndvi"):
                story.append(Spacer(1, 0.08 * inch))
                story.append(Paragraph(
                    f"<b>Imagen NDVI — Índice de Vegetación ({year_str})</b>",
                    self.styles["BodySmall"],
                ))
                try:
                    ndvi_img = io.BytesIO(year_thumb["ndvi"])
                    story.append(RLImage(ndvi_img, width=4.5 * inch, height=3 * inch))
                except Exception:
                    pass

            # ── Expanded stats table ──
            def_stats = yd.get("deforestation", {}).get("stats", {})
            ue_stats = yd.get("urban_expansion", {}).get("stats", {})
            conf = _to_float(def_stats.get("confidence"))
            conf_str = f"{conf * 100:.0f}%" if conf else "—"

            mini_data = [
                ["Deforestación", f"{_to_float(def_stats.get('area_ha')):.2f} ha",
                 f"{int(_to_float(def_stats.get('n_features')))} zona(s)"],
                ["Confianza deforestación", conf_str, ""],
                ["Exp. urbana", f"{_to_float(ue_stats.get('area_ha')):.2f} ha",
                 f"{int(_to_float(ue_stats.get('n_features')))} zona(s)"],
            ]
            # Add all vegetation classes with > 0.5%
            for cls_key in VEG_COLORS:
                val = _to_float(classes.get(cls_key))
                if val > 0.5:
                    mini_data.append([
                        VEG_LABELS.get(cls_key, cls_key),
                        f"{val:.1f}%",
                        "",
                    ])

            # Alertas from features
            feats = yd.get("urban_expansion", {}).get("geojson", {}).get("features", [])
            alertas = [f.get("properties", {}).get("alerta", "")
                       for f in feats if f.get("properties", {}).get("alerta")]
            if alertas:
                mini_data.append(["Alertas", "; ".join(set(alertas)), "⚠"])

            mini_t = Table(
                [["Indicador", "Valor", "Detalle"]] + mini_data,
                colWidths=[2.2 * inch, 2.2 * inch, 2.1 * inch],
            )
            mini_t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#374151")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#e5e7eb")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1),
                 [colors.white, colors.HexColor("#f9fafb")]),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            story.append(mini_t)

            # Page break between years (except the last)
            if year_str != sorted(timeline.keys())[-1]:
                story.append(PageBreak())

        return story

    # ── 6. Conclusions / Recommendations ──
    def _build_conclusions(self, interp: dict) -> list:
        story: list = []
        story.append(PageBreak())
        story.append(Paragraph("6. CONCLUSIONES Y RECOMENDACIONES", self.styles["SectionHead"]))

        story.append(Paragraph(
            f"<b>Nivel de riesgo ambiental:</b> {interp['level']} ({interp['score']}/100)", self.styles["Body"]))
        story.append(Paragraph(f"<b>Tendencia:</b> {interp['trend']}", self.styles["Body"]))

        story.append(Spacer(1, 0.1 * inch))
        story.append(Paragraph("<b>Recomendaciones:</b>", self.styles["Body"]))
        for i, rec in enumerate(interp["recommendations"], 1):
            story.append(Paragraph(f"{i}. {rec}", self.styles["BodySmall"]))
        return story

    # ── 7. Methodology ──
    def _build_methodology(self) -> list:
        story: list = []
        story.append(Spacer(1, 0.3 * inch))
        story.append(Paragraph("7. ANEXO METODOLÓGICO", self.styles["SectionHead"]))

        texts = [
            "Las imágenes fueron obtenidas del conjunto de datos <b>Dynamic World</b> (Google/WRI) "
            "con resolución espacial de 10 metros, derivado de Sentinel-2. Se procesaron con los "
            "motores de inteligencia artificial del sistema APEX.",
            "<b>Motor de Deforestación:</b> Compara la clasificación de cobertura terrestre entre "
            "dos períodos (T1 vs T2). Se identifican transiciones de clases forestales (bosque denso, "
            "manglar, matorral) hacia clases no forestales (pastizal, suelo, urbano).",
            "<b>Motor de Expansión Urbana:</b> Detecta transiciones hacia la clase 'urbano' desde "
            "cualquier otra cobertura. Incluye alertas automáticas para posibles construcciones "
            "sin permiso.",
            "<b>Motor de Vegetación:</b> Clasificación de 11 clases de cobertura terrestre basada "
            "en Dynamic World V1. Se calcula el porcentaje de cada clase sobre el AOI.",
            "<b>Hansen Global Forest Change (v1.12):</b> Datos históricos de pérdida forestal del "
            "dataset UMD/Hansen a resolución 30m. Cubre el periodo 2001-2024 y utiliza bandas de "
            "treecover2000, loss y lossyear para vectorizar zonas de pérdida con filtro de bosque ≥30%.",
            "<b>Alertas GLAD-S2 / RADD:</b> GLAD-S2 utiliza Sentinel-2 (óptico) para detección de "
            "deforestación confirmada. RADD utiliza Sentinel-1 (SAR) que penetra nubes — crítico para "
            "regiones tropicales como Chiapas, Oaxaca y Veracruz.",
            "<b>Drivers de Deforestación (WRI):</b> Clasificación de causas de deforestación a 1km "
            "desarrollada por WRI / Google DeepMind. Incluye: agricultura permanente, commodities, "
            "cultivo rotacional, tala, incendios, asentamientos e infraestructura.",
            "<b>Incendios (MODIS MCD64A1):</b> Áreas quemadas detectadas por MODIS a resolución 500m. "
            "Se correlaciona con polígonos de deforestación para identificar causalidad.",
            "<b>Puntos de Calor FIRMS:</b> Detecciones de fuego activo en tiempo casi real del sistema "
            "NASA FIRMS con sensores VIIRS (375 m, satélites S-NPP y NOAA-20) y MODIS (1 km, Terra y Aqua). "
            "Se reporta FRP (Fire Radiative Power) en megawatts como indicador de intensidad. "
            "Datos locales indexados desde 2018 con respaldo de API en línea.",
            "<b>Contexto Legal (WDPA):</b> Intersección del AOI con Áreas Naturales Protegidas (ANPs) "
            "de la CONANP/WDPA. La deforestación dentro de una ANP constituye infracción potencial "
            "bajo el Artículo 47 de la LGEEPA.",
            "<b>Detección de Anomalías:</b> Se calcula el z-score de cada indicador a lo largo de la "
            "serie temporal. Valores con z &gt; 2σ se clasifican como anomalías de severidad 'media' "
            "(2-3σ) o 'alta' (&gt;3σ).",
            "<b>Nivel de Riesgo:</b> Score compuesto (0-100) que pondera: deforestación acumulada, "
            "expansión urbana, cambio en cobertura forestal, número de anomalías y tasa anual.",
        ]
        for t in texts:
            story.append(Paragraph(t, self.styles["BodySmall"]))
        return story

    # ── PUBLIC API ──
    def generate(self, summary: dict, job_id: str,
                 aoi_geojson: dict = None,
                 thumbnails: dict = None,
                 overview_thumb: dict = None) -> io.BytesIO:
        """Generate full institutional PDF report.

        Args:
            summary:        timeline_summary from DB
            job_id:         analysis job ID
            aoi_geojson:    original AOI polygon (for map overlays)
            thumbnails:     {year: {rgb: bytes, ndvi: bytes, bounds: [...]}}
            overview_thumb: {rgb: bytes, bounds: [...]} for location map
        """
        folio = f"PROFEPA-APEX-{job_id[:8].upper()}"
        cumulative = summary.get("cumulative", {})
        anomalies = summary.get("anomalies", [])
        timeline = summary.get("timeline", {})
        period = cumulative.get("period", "N/D")
        interp = _generate_interpretation(timeline, cumulative, anomalies)
        thumbnails = thumbnails or {}

        # Compute AOI coordinates + area
        _coords = self._compute_coords(aoi_geojson)
        _area_ha = _to_float(cumulative.get("total_area_ha",
                     cumulative.get("area_ha", 0)))

        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf, pagesize=letter,
            topMargin=1 * inch, bottomMargin=0.8 * inch,
            leftMargin=0.75 * inch, rightMargin=0.75 * inch,
        )

        story: list = []
        story += self._build_cover(folio, period, interp)

        # ── NEW: Location overview map ──
        if aoi_geojson:
            story += self._build_location_section(aoi_geojson, _coords,
                                                  _area_ha, overview_thumb)

        # ── NEW: Geographic info section ──
        story += self._build_geographic_info(cumulative, _coords, _area_ha)

        story += self._build_summary(cumulative, len(anomalies), interp)
        story += self._build_legal_context(summary)
        story += self._build_hansen_section(summary)
        story += self._build_drivers_section(summary)
        story += self._build_fire_section(summary)
        story += self._build_deforestation_alerts_section(summary)
        story += self._build_sar_section(summary)
        story += self._build_crossval_section(summary)
        story += self._build_firms_section(summary)
        story += self._build_alerts(anomalies)
        if timeline:
            story += self._build_annual_data(timeline)
            # ── NEW: Full vegetation breakdown table ──
            story += self._build_vegetation_table(timeline)
            story += self._build_charts(timeline)
            story += self._build_year_pages(timeline, interp,
                                            aoi_geojson=aoi_geojson,
                                            thumbnails=thumbnails)
        story += self._build_conclusions(interp)
        story += self._build_methodology()

        story.append(Spacer(1, 0.3 * inch))
        story.append(Paragraph(
            "— Fin del reporte — Generado automáticamente por el sistema APEX —",
            self.styles["Small"],
        ))

        doc.build(story, onFirstPage=self._header_footer, onLaterPages=self._header_footer)
        buf.seek(0)
        return buf

    # ── Helper: compute coords from AOI ──
    @staticmethod
    def _compute_coords(aoi_geojson: dict) -> dict:
        if not aoi_geojson:
            return {}
        try:
            gdf = _geojson_to_gdf(aoi_geojson)
            if gdf is None or gdf.empty:
                return {}
            minx, miny, maxx, maxy = gdf.total_bounds
            return {
                "center_lat": (miny + maxy) / 2,
                "center_lon": (minx + maxx) / 2,
                "north": maxy, "south": miny,
                "east": maxx, "west": minx,
            }
        except Exception:
            return {}

    # ── NEW: Location overview section ──
    def _build_location_section(self, aoi_geojson, coords, area_ha, overview_thumb):
        story: list = []
        story.append(Paragraph("UBICACIÓN GEOGRÁFICA", self.styles["SectionHead"]))
        try:
            loc_buf = _build_location_map(aoi_geojson, coords, area_ha, overview_thumb)
            story.append(RLImage(loc_buf, width=6.5 * inch, height=5 * inch))
        except Exception:
            story.append(Paragraph(
                "No se pudo generar el mapa de ubicación.",
                self.styles["BodySmall"],
            ))
        story.append(PageBreak())
        return story

    # ── NEW: Geographic info table ──
    def _build_geographic_info(self, cumulative, coords, area_ha):
        story: list = []
        if not coords:
            return story
        story.append(Paragraph("INFORMACIÓN GEOGRÁFICA", self.styles["SectionHead"]))
        data = [
            ["Parámetro", "Valor"],
            ["Latitud centro", f'{coords.get("center_lat", 0):.6f}°'],
            ["Longitud centro", f'{coords.get("center_lon", 0):.6f}°'],
            ["Límite norte", f'{coords.get("north", 0):.6f}°'],
            ["Límite sur", f'{coords.get("south", 0):.6f}°'],
            ["Límite este", f'{coords.get("east", 0):.6f}°'],
            ["Límite oeste", f'{coords.get("west", 0):.6f}°'],
        ]
        if area_ha > 0:
            data.append(["Área total", f"{area_ha:.2f} ha"])
        period = cumulative.get("period", "N/D")
        data.append(["Período analizado", period])
        n_years = cumulative.get("years_analyzed", "N/D")
        data.append(["Años analizados", str(n_years)])

        t = Table(data, colWidths=[3 * inch, 3.5 * inch])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), self.COLOR_SECONDARY),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d1d5db")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f4f8")]),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(t)
        story.append(Spacer(1, 0.15 * inch))
        return story

    # ── NEW: Full vegetation breakdown table ──
    def _build_vegetation_table(self, timeline: dict) -> list:
        """Table showing all vegetation classes per year (not just bosque + urbano)."""
        story: list = []
        story.append(Spacer(1, 0.2 * inch))
        story.append(Paragraph("COMPOSICIÓN DE VEGETACIÓN POR AÑO (%)", self.styles["SectionHead"]))

        years = sorted(timeline.keys())
        # Gather all classes that appear
        all_classes = set()
        for y in years:
            classes = timeline[y].get("vegetation", {}).get("stats", {}).get("classes", {})
            all_classes.update(k for k, v in classes.items() if _to_float(v) > 0.1)
        ordered = [c for c in VEG_COLORS if c in all_classes]
        if not ordered:
            return story

        header = ["Año"] + [VEG_LABELS.get(c, c) for c in ordered]
        rows = [header]
        for y in years:
            classes = timeline[y].get("vegetation", {}).get("stats", {}).get("classes", {})
            row = [y] + [f"{_to_float(classes.get(c)):.1f}" for c in ordered]
            rows.append(row)

        col_w = min(1.2 * inch, 6.5 * inch / (len(ordered) + 1))
        col_widths = [0.7 * inch] + [col_w] * len(ordered)

        t = Table(rows, colWidths=col_widths)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#166534")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 7),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#d1d5db")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0fdf4")]),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ALIGN", (1, 1), (-1, -1), "CENTER"),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        story.append(t)
        story.append(Spacer(1, 0.05 * inch))
        story.append(Paragraph(
            "Valores expresados como porcentaje del área total de interés. "
            "Fuente: Dynamic World (Google/WRI, 10 m).",
            self.styles["Small"],
        ))
        return story


# ═══════════════════════════════════════════════════════════
#  Word Report Generator
# ═══════════════════════════════════════════════════════════

class APEXWordReportGenerator:
    """Genera reportes Word profesionales completos para APEX / PROFEPA."""

    COLOR_PRIMARY = RGBColor(0, 212, 255)
    COLOR_SECONDARY = RGBColor(30, 58, 95)
    COLOR_RED = RGBColor(239, 68, 68)
    COLOR_GREEN = RGBColor(46, 160, 67)

    def __init__(self):
        plt.rcParams.update({
            "font.family": ["DejaVu Sans", "Arial", "sans-serif"],
            "font.size": 9,
            "axes.grid": True,
            "grid.alpha": 0.3,
            "figure.facecolor": "white",
        })

    def _add_styles(self, doc: Document):
        style = doc.styles.add_style("APEXTitle", 1)
        fmt = style.font
        fmt.name, fmt.size, fmt.bold, fmt.color.rgb = "Arial", Pt(20), True, self.COLOR_SECONDARY

        style2 = doc.styles.add_style("APEXHeading", 1)
        fmt2 = style2.font
        fmt2.name, fmt2.size, fmt2.bold, fmt2.color.rgb = "Arial", Pt(14), True, self.COLOR_PRIMARY

        style3 = doc.styles.add_style("APEXSubhead", 1)
        fmt3 = style3.font
        fmt3.name, fmt3.size, fmt3.bold, fmt3.color.rgb = "Arial", Pt(11), True, RGBColor(55, 65, 81)

        style4 = doc.styles.add_style("APEXBody", 1)
        fmt4 = style4.font
        fmt4.name, fmt4.size, fmt4.color.rgb = "Arial", Pt(10), RGBColor(31, 41, 55)

        style5 = doc.styles.add_style("APEXSmall", 1)
        fmt5 = style5.font
        fmt5.name, fmt5.size, fmt5.color.rgb = "Arial", Pt(9), RGBColor(75, 85, 99)

    def _add_cover(self, doc: Document, folio: str, period: str, interp: dict):
        for _ in range(4):
            doc.add_paragraph("")
        p = doc.add_paragraph("REPORTE DE ANÁLISIS\nPREDICTIVO ECOSISTÉMICO", style="APEXTitle")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph("")

        # Risk level
        risk_p = doc.add_paragraph(style="APEXHeading")
        risk_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = risk_p.add_run(f"NIVEL DE RIESGO: {interp['level']}  ({interp['score']}/100)")
        run.bold = True
        run.font.size = Pt(14)

        doc.add_paragraph("")
        meta = doc.add_paragraph(style="APEXBody")
        meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for line in [
            f"Folio: {folio}",
            f"Período: {period}",
            f"Fecha: {datetime.now().strftime('%Y-%m-%d %H:%M UTC')}",
            "Sistema: APEX — Análisis Predictivo de Ecosistemas con IA",
            "Institución: PROFEPA",
        ]:
            run = meta.add_run(line + "\n")
            run.font.size = Pt(11)
        doc.add_page_break()

    def _add_summary(self, doc: Document, cumulative: dict, n_anomalies: int, interp: dict):
        doc.add_paragraph("1. Resumen Ejecutivo", style="APEXHeading")

        rows_data = [
            ("Indicador", "Valor", "Interpretación"),
            ("Deforestación acumulada", f"{cumulative.get('total_deforestation_ha', 0)} ha",
             "Alta" if _to_float(cumulative.get("total_deforestation_ha")) > 10 else "Normal"),
            ("Expansión urbana total", f"{cumulative.get('total_urban_expansion_ha', 0)} ha",
             "Alta" if _to_float(cumulative.get("total_urban_expansion_ha")) > 5 else "Normal"),
            ("Cambio bosque denso", f"{cumulative.get('bosque_denso_change_pct', 0)} pp",
             "Pérdida" if _to_float(cumulative.get("bosque_denso_change_pct")) < -1 else "Estable"),
            ("Cambio urbano", f"{cumulative.get('urbano_change_pct', 0)} pp",
             "Creciendo" if _to_float(cumulative.get("urbano_change_pct")) > 1 else "Estable"),
            ("Años analizados", str(cumulative.get("years_analyzed", "-")), ""),
            ("Anomalías", str(n_anomalies),
             "⚠ Requiere atención" if n_anomalies > 0 else "Sin alertas"),
        ]
        table = doc.add_table(rows=len(rows_data), cols=3)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (a, b, c) in enumerate(rows_data):
            table.rows[i].cells[0].text = a
            table.rows[i].cells[1].text = b
            table.rows[i].cells[2].text = c
            if i == 0:
                for cell in table.rows[i].cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
        doc.add_paragraph("")
        trend_p = doc.add_paragraph(style="APEXBody")
        run = trend_p.add_run("Tendencia general: ")
        run.bold = True
        trend_p.add_run(interp["trend"])
        doc.add_paragraph("")

    def _add_legal_context(self, doc: Document, summary: dict):
        """Sección de contexto legal: intersección con ANPs (WDPA) — Word."""
        lc = summary.get("legal_context")
        if not lc:
            return

        doc.add_paragraph("Contexto Legal — Áreas Naturales Protegidas", style="APEXHeading")

        anp_name = lc.get("anp_name", "Desconocida")
        category = lc.get("category", "N/D")
        overlap_area = _to_float(lc.get("overlap_area_ha"))
        overlap_pct = _to_float(lc.get("overlap_pct"))
        risk_assessment = lc.get("risk_assessment", "No disponible")
        intersects = lc.get("intersects_anp", False)

        rows_data = [
            ("Parámetro", "Valor"),
            ("Nombre del ANP", anp_name),
            ("Categoría", category),
            ("Área de traslape", f"{overlap_area:.2f} ha"),
            ("Porcentaje de traslape", f"{overlap_pct:.1f}%"),
            ("Evaluación de riesgo", risk_assessment),
        ]
        table = doc.add_table(rows=len(rows_data), cols=2)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (a, b) in enumerate(rows_data):
            table.rows[i].cells[0].text = a
            table.rows[i].cells[1].text = b
            if i == 0:
                for cell in table.rows[i].cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

        if intersects:
            doc.add_paragraph("")
            warn_p = doc.add_paragraph(style="APEXBody")
            run = warn_p.add_run(
                "ADVERTENCIA: El área de interés se traslapa con un Área Natural Protegida. "
                "La deforestación dentro de una ANP puede constituir una infracción conforme al "
                "Artículo 47 de la LGEEPA (Ley General del Equilibrio Ecológico y la "
                "Protección al Ambiente). Se recomienda verificación en campo y, en su caso, "
                "iniciar procedimiento administrativo."
            )
            run.bold = True
            run.font.color.rgb = RGBColor(239, 68, 68)
        doc.add_paragraph("")

    def _add_hansen_section(self, doc: Document, summary: dict):
        """Sección Hansen: referencia histórica de pérdida forestal (2001-2024) — Word."""
        hansen = summary.get("hansen")
        if not hansen:
            return

        doc.add_paragraph("Hansen — Referencia Histórica (2001-2024)", style="APEXHeading")

        total_loss = _to_float(hansen.get("total_loss_ha"))
        n_features = hansen.get("n_features", 0)
        avg_treecover = _to_float(hansen.get("avg_treecover_pct"))

        intro_p = doc.add_paragraph(style="APEXBody")
        intro_p.add_run(
            f"Según el dataset Hansen Global Forest Change v1.12 (UMD), se registra una "
            f"pérdida forestal acumulada de {total_loss:.2f} ha distribuida en "
            f"{n_features} polígono(s), con una cobertura arbórea promedio inicial de "
            f"{avg_treecover:.1f}%."
        )

        # Loss by year table
        loss_by_year = hansen.get("loss_by_year", {})
        if loss_by_year:
            doc.add_paragraph("Desglose de pérdida por año:", style="APEXSubhead")
            sorted_years = sorted(loss_by_year.keys())
            table = doc.add_table(rows=1 + len(sorted_years), cols=2)
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.rows[0].cells[0].text = "Año"
            table.rows[0].cells[1].text = "Pérdida (ha)"
            for cell in table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            for j, yr in enumerate(sorted_years):
                table.rows[j + 1].cells[0].text = str(yr)
                table.rows[j + 1].cells[1].text = f"{_to_float(loss_by_year[yr]):.2f}"

        source_p = doc.add_paragraph(style="APEXSmall")
        source_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        source_p.add_run("Fuente: UMD / Hansen — Global Forest Change v1.12 (resolución 30 m)")
        doc.add_paragraph("")

    def _add_drivers_section(self, doc: Document, summary: dict):
        """Sección de drivers de deforestación (WRI / Google DeepMind) — Word."""
        drivers = summary.get("drivers")
        if not drivers:
            return

        doc.add_paragraph("Clasificación de Drivers de Deforestación", style="APEXHeading")
        doc.add_paragraph(
            "Distribución de las causas de deforestación identificadas en el área de interés, "
            "según la clasificación de WRI / Google DeepMind:",
            style="APEXBody",
        )

        # Build distribution table
        distribution = drivers.get("distribution", drivers.get("classes",
                       drivers.get("drivers", {})))
        parsed_rows = []
        if isinstance(distribution, list):
            for item in distribution:
                name = item.get("driver", item.get("name", "Desconocido"))
                pct = _to_float(item.get("percentage", item.get("pct")))
                parsed_rows.append((name, f"{pct:.1f}%"))
        elif isinstance(distribution, dict):
            for driver_name in sorted(distribution.keys()):
                pct = _to_float(distribution[driver_name])
                parsed_rows.append((driver_name, f"{pct:.1f}%"))

        if parsed_rows:
            table = doc.add_table(rows=1 + len(parsed_rows), cols=2)
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.rows[0].cells[0].text = "Driver"
            table.rows[0].cells[1].text = "Porcentaje"
            for cell in table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            for j, (name, pct_str) in enumerate(parsed_rows):
                table.rows[j + 1].cells[0].text = name
                table.rows[j + 1].cells[1].text = pct_str

        source_p = doc.add_paragraph(style="APEXSmall")
        source_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        source_p.add_run("Fuente: WRI / Google DeepMind — Drivers de Deforestación (resolución 1 km)")
        doc.add_paragraph("")

    def _add_fire_section(self, doc: Document, summary: dict):
        """Sección de incendios / áreas quemadas (MODIS) — Word."""
        fire = summary.get("fire")
        if not fire:
            return

        doc.add_paragraph("Incendios y Áreas Quemadas (MODIS)", style="APEXHeading")

        total_burned = _to_float(fire.get("total_burned_ha"))
        fire_count = fire.get("fire_count", 0)
        fire_defor_pct = _to_float(fire.get("fire_related_deforestation_pct"))

        p = doc.add_paragraph(style="APEXBody")
        p.add_run(
            f"Se detectaron {fire_count} zona(s) quemada(s) con un área total de "
            f"{total_burned:.2f} ha según el sensor MODIS MCD64A1 (resolución 500 m). "
            f"Se estima que el {fire_defor_pct:.1f}% de la deforestación detectada "
            f"está asociada espacialmente con áreas quemadas."
        )

        source_p = doc.add_paragraph(style="APEXSmall")
        source_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        source_p.add_run("Fuente: MODIS MCD64A1 — Burned Area (resolución 500 m)")
        doc.add_paragraph("")

    def _add_deforestation_alerts_section(self, doc: Document, summary: dict):
        """Sección de alertas GLAD/RADD — Word."""
        alerts = summary.get("alerts")
        if not alerts:
            return

        doc.add_paragraph("Alertas de Deforestación (GLAD / RADD)", style="APEXHeading")

        total_alerts = alerts.get("total_alerts", 0)
        glad_count = alerts.get("glad_count", 0)
        radd_count = alerts.get("radd_count", 0)
        confirmed = alerts.get("confirmed_count", 0)
        total_area = _to_float(alerts.get("total_area_ha"))

        p = doc.add_paragraph(style="APEXBody")
        p.add_run(
            f"Se identificaron {total_alerts} alertas de deforestación activas: "
            f"{glad_count} GLAD (Sentinel-2 óptico) y {radd_count} RADD "
            f"(Sentinel-1 SAR, penetra nubes), cubriendo un área total de {total_area:.2f} ha. "
            f"De estas, {confirmed} tienen confianza alta/confirmada."
        )

        if radd_count > 0 and glad_count == 0:
            note = doc.add_paragraph(style="APEXBody")
            run = note.add_run(
                "Nota: Solo se detectaron alertas SAR (RADD), lo que sugiere "
                "posible cobertura de nubes persistente que impide la detección óptica."
            )
            run.italic = True

        source_p = doc.add_paragraph(style="APEXSmall")
        source_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        source_p.add_run("Fuente: GLAD-S2 (UMD, 10 m óptico) + RADD (WUR, 10 m SAR)")
        doc.add_paragraph("")

    def _add_sar_section(self, doc: Document, summary: dict):
        """Sección de detección de cambios SAR — Word."""
        sar = summary.get("sar")
        if not sar:
            return

        doc.add_paragraph("Detección de Cambios SAR (Sentinel-1)", style="APEXHeading")

        n_changes = sar.get("sar_change_count", 0)
        total_area = _to_float(sar.get("total_change_ha", sar.get("total_area_ha")))
        high_conf = sar.get("high_confidence_count", 0)

        p = doc.add_paragraph(style="APEXBody")
        p.add_run(
            f"Mediante análisis bi-temporal de log-ratio sobre composites radar "
            f"Sentinel-1 (banda C, VV+VH), se detectaron {n_changes} "
            f"zonas de cambio con un área total de {total_area:.2f} ha. "
            f"De estas, {high_conf} presentan alta confianza (caída >5 dB). "
            f"La detección SAR es complementaria a la óptica y funciona al 100% bajo "
            f"cobertura de nubes."
        )

        source_p = doc.add_paragraph(style="APEXSmall")
        source_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        source_p.add_run("Fuente: Sentinel-1 GRD IW — Log-ratio change detection (resolución 10 m)")
        doc.add_paragraph("")

    def _add_crossval_section(self, doc: Document, summary: dict):
        """Sección de validación cruzada DW vs MapBiomas — Word."""
        cv = summary.get("crossval")
        if not cv:
            return

        doc.add_paragraph("Validación Cruzada (Dynamic World vs MapBiomas)", style="APEXHeading")

        agreement = _to_float(cv.get("agreement_pct"))
        disagreement = cv.get("disagreement_zones", 0)
        total = cv.get("total_compared", 0)

        p = doc.add_paragraph(style="APEXBody")
        p.add_run(
            f"Se compararon {total} detecciones de deforestación (Dynamic World) contra "
            f"la clasificación LULC de MapBiomas México. La tasa de acuerdo es del "
            f"{agreement:.1f}%, con {disagreement} zona(s) de desacuerdo "
            f"identificadas como posibles falsos positivos."
        )

        if agreement >= 80:
            note = doc.add_paragraph(style="APEXBody")
            note.add_run(
                "La alta tasa de acuerdo refuerza la confiabilidad de las detecciones."
            )
        elif agreement < 60:
            note = doc.add_paragraph(style="APEXBody")
            run = note.add_run(
                "Nota: La tasa de acuerdo moderada sugiere prudencia en la "
                "interpretación. Se recomienda verificación en campo."
            )
            run.bold = True
            run.font.color.rgb = RGBColor(239, 68, 68)

        source_p = doc.add_paragraph(style="APEXSmall")
        source_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        source_p.add_run("Fuente: MapBiomas México v1.0 (resolución 30 m)")
        doc.add_paragraph("")

    def _add_firms_section(self, doc: Document, summary: dict):
        """Sección de puntos de calor FIRMS — Word."""
        timeline = summary.get("timeline", {})
        cumulative = summary.get("cumulative", {})
        total_hotspots = cumulative.get("total_firms_hotspots", 0)
        total_frp = _to_float(cumulative.get("total_frp_mw", 0))

        if total_hotspots == 0:
            return

        doc.add_paragraph("Puntos de Calor FIRMS (VIIRS / MODIS NRT)", style="APEXHeading")

        p = doc.add_paragraph(style="APEXBody")
        p.add_run(
            f"Se identificaron {total_hotspots} detecciones de puntos de calor activos "
            f"en el área de estudio durante el período analizado, con una potencia radiativa "
            f"total de {total_frp:,.1f} MW (Fire Radiative Power). "
            f"Los datos provienen del sistema NASA FIRMS con sensores VIIRS (375 m) y MODIS (1 km)."
        )

        # Per-year table
        headers = ["Año", "Hotspots", "Alta conf.", "Clusters", "FRP total (MW)", "FRP máx (MW)", "Satélites"]
        rows_data = []
        for yr in sorted(timeline.keys()):
            yr_firms = timeline[yr].get("firms_hotspots", {}).get("stats", {})
            hc = yr_firms.get("hotspot_count", 0)
            if hc == 0:
                continue
            hi = yr_firms.get("high_confidence_count", 0)
            cl = yr_firms.get("cluster_count", 0)
            frp_t = yr_firms.get("total_frp_mw", 0)
            frp_m = yr_firms.get("max_frp_mw", 0)
            sats = ", ".join(yr_firms.get("satellites", []))
            rows_data.append([str(yr), str(hc), str(hi), str(cl),
                              f"{frp_t:,.1f}", f"{frp_m:,.1f}", sats])

        if rows_data:
            t = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
            t.style = "Light Grid Accent 1"
            for ci, h in enumerate(headers):
                t.rows[0].cells[ci].text = h
                for paragraph in t.rows[0].cells[ci].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            for ri, row in enumerate(rows_data, 1):
                for ci, val in enumerate(row):
                    t.rows[ri].cells[ci].text = val

        doc.add_paragraph("")
        note = doc.add_paragraph(style="APEXSmall")
        run = note.add_run(
            "FRP (Fire Radiative Power): Potencia radiativa del fuego medida en megawatts. "
            "Valores altos indican mayor intensidad de quema y potencial de daño."
        )
        run.italic = True
        source_p = doc.add_paragraph(style="APEXSmall")
        source_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        source_p.add_run(
            "Fuente: NASA FIRMS — VIIRS (S-NPP, NOAA-20) 375 m / MODIS (Terra, Aqua) 1 km, datos NRT + archivo"
        )
        doc.add_paragraph("")

    def _add_alerts(self, doc: Document, anomalies: list):
        doc.add_paragraph("2. Alertas y Anomalías", style="APEXHeading")
        if not anomalies:
            doc.add_paragraph(
                "No se detectaron anomalías estadísticas (z-score > 2σ) en la serie temporal analizada. "
                "Todos los indicadores se encuentran dentro de la variabilidad histórica esperada.",
                style="APEXBody",
            )
            return
        doc.add_paragraph(
            f"Se detectaron {len(anomalies)} anomalía(s) estadística(s):", style="APEXBody")
        headers = ["Año", "Motor", "Valor (ha)", "Media", "z-score", "Sev."]
        table = doc.add_table(rows=1 + len(anomalies), cols=6)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            for run in table.rows[0].cells[i].paragraphs[0].runs:
                run.bold = True
        for j, a in enumerate(anomalies):
            table.rows[j + 1].cells[0].text = str(a.get("year", ""))
            table.rows[j + 1].cells[1].text = a.get("engine", "").replace("_", " ").title()
            table.rows[j + 1].cells[2].text = f"{_to_float(a.get('value', a.get('area_ha'))):.1f}"
            table.rows[j + 1].cells[3].text = f"{_to_float(a.get('mean_ha')):.1f}"
            table.rows[j + 1].cells[4].text = f"{_to_float(a.get('z_score')):.2f}"
            table.rows[j + 1].cells[5].text = a.get("severity", "").upper()
        # Anomaly messages
        for a in anomalies:
            msg = a.get("message", "")
            if msg:
                doc.add_paragraph(f"• {msg}", style="APEXSmall")
        doc.add_paragraph("")

    def _add_annual_data(self, doc: Document, timeline: dict):
        doc.add_paragraph("3. Datos Tabulares por Año", style="APEXHeading")
        cols = ["Año", "Base", "Def. (ha)", "Exp.Urb. (ha)", "Bosque (%)", "Urb. (%)", "Conf."]
        years = sorted(timeline.keys())
        table = doc.add_table(rows=1 + len(years), cols=7)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, c in enumerate(cols):
            table.rows[0].cells[i].text = c
            for run in table.rows[0].cells[i].paragraphs[0].runs:
                run.bold = True
        for j, y in enumerate(years):
            yd = timeline[y]
            def_stats = yd.get("deforestation", {}).get("stats", {})
            ue_stats = yd.get("urban_expansion", {}).get("stats", {})
            classes = yd.get("vegetation", {}).get("stats", {}).get("classes", {})
            conf = _to_float(def_stats.get("confidence"))
            table.rows[j + 1].cells[0].text = y
            table.rows[j + 1].cells[1].text = str(yd.get("baseline_year", "?"))
            table.rows[j + 1].cells[2].text = f"{_to_float(def_stats.get('area_ha')):.1f}"
            table.rows[j + 1].cells[3].text = f"{_to_float(ue_stats.get('area_ha')):.1f}"
            table.rows[j + 1].cells[4].text = f"{_to_float(classes.get('bosque_denso')):.1f}"
            table.rows[j + 1].cells[5].text = f"{_to_float(classes.get('urbano')):.1f}"
            table.rows[j + 1].cells[6].text = f"{conf*100:.0f}%" if conf else "—"
        doc.add_paragraph("")

    def _add_charts(self, doc: Document, timeline: dict):
        doc.add_page_break()
        doc.add_paragraph("4. Gráficas de Evolución Temporal", style="APEXHeading")

        doc.add_paragraph("4.1 Cambios detectados por año", style="APEXSubhead")
        buf1 = _build_change_bar_chart(timeline)
        doc.add_picture(buf1, width=Inches(5.8))
        doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph("")

        doc.add_paragraph("4.2 Evolución de la composición de vegetación", style="APEXSubhead")
        buf2 = _build_veg_stacked_area(timeline)
        doc.add_picture(buf2, width=Inches(5.8))
        doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def _add_year_pages(self, doc: Document, timeline: dict, interp: dict,
                       aoi_geojson: dict = None, thumbnails: dict = None):
        thumbnails = thumbnails or {}
        doc.add_page_break()
        doc.add_paragraph("5. Análisis Detallado por Año", style="APEXHeading")

        for year_str in sorted(timeline.keys()):
            yd = timeline[year_str]
            baseline = yd.get("baseline_year", int(year_str) - 1)
            doc.add_paragraph(f"Año {year_str} (base: {baseline})", style="APEXSubhead")

            # Interpretation
            txt = interp["year_interpretations"].get(year_str, "")
            if txt:
                doc.add_paragraph(txt, style="APEXSmall")

            # Cartographic Map
            year_thumb = thumbnails.get(year_str)
            map_buf = _build_year_map(yd, year_str,
                                      aoi_geojson=aoi_geojson,
                                      year_thumb=year_thumb)
            doc.add_picture(map_buf, width=Inches(4.5))
            doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # NDVI satellite image (if available)
            if year_thumb and year_thumb.get("ndvi"):
                doc.add_paragraph(
                    f"Imagen NDVI — Índice de Vegetación ({year_str})",
                    style="APEXSubhead",
                )
                try:
                    ndvi_buf = io.BytesIO(year_thumb["ndvi"])
                    doc.add_picture(ndvi_buf, width=Inches(4))
                    doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                except Exception:
                    pass

            # Donut
            classes = yd.get("vegetation", {}).get("stats", {}).get("classes", {})
            donut_buf = _build_veg_donut(classes, year_str)
            doc.add_picture(donut_buf, width=Inches(3.5))
            doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Expanded stats table
            def_stats = yd.get("deforestation", {}).get("stats", {})
            ue_stats = yd.get("urban_expansion", {}).get("stats", {})
            conf = _to_float(def_stats.get("confidence"))
            conf_str = f"{conf * 100:.0f}%" if conf else "—"

            mini_rows = [
                ("Indicador", "Valor", "Detalle"),
                ("Deforestación", f"{_to_float(def_stats.get('area_ha')):.2f} ha",
                 f"{int(_to_float(def_stats.get('n_features')))} zona(s)"),
                ("Confianza", conf_str, ""),
                ("Exp. urbana", f"{_to_float(ue_stats.get('area_ha')):.2f} ha",
                 f"{int(_to_float(ue_stats.get('n_features')))} zona(s)"),
            ]
            # All vegetation classes > 0.5%
            for cls_key in VEG_COLORS:
                val = _to_float(classes.get(cls_key))
                if val > 0.5:
                    mini_rows.append((VEG_LABELS.get(cls_key, cls_key), f"{val:.1f}%", ""))

            t = doc.add_table(rows=len(mini_rows), cols=3)
            t.style = "Light Grid Accent 1"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, (a, b, c) in enumerate(mini_rows):
                t.rows[i].cells[0].text = a
                t.rows[i].cells[1].text = b
                t.rows[i].cells[2].text = c
                if i == 0:
                    for cell in t.rows[i].cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
            doc.add_paragraph("")

            if year_str != sorted(timeline.keys())[-1]:
                doc.add_page_break()

    def _add_conclusions(self, doc: Document, interp: dict):
        doc.add_page_break()
        doc.add_paragraph("6. Conclusiones y Recomendaciones", style="APEXHeading")

        p = doc.add_paragraph(style="APEXBody")
        run = p.add_run(f"Nivel de riesgo ambiental: {interp['level']} ({interp['score']}/100)")
        run.bold = True
        p2 = doc.add_paragraph(style="APEXBody")
        run2 = p2.add_run("Tendencia: ")
        run2.bold = True
        p2.add_run(interp["trend"])
        doc.add_paragraph("")

        doc.add_paragraph("Recomendaciones:", style="APEXSubhead")
        for i, rec in enumerate(interp["recommendations"], 1):
            doc.add_paragraph(f"{i}. {rec}", style="APEXSmall")

    def _add_methodology(self, doc: Document):
        doc.add_page_break()
        doc.add_paragraph("7. Anexo Metodológico", style="APEXHeading")

        texts = [
            "Las imágenes fueron obtenidas del conjunto de datos Dynamic World (Google/WRI) "
            "con resolución espacial de 10 metros, derivado de Sentinel-2. Se procesaron con los "
            "motores de inteligencia artificial del sistema APEX.",
            "Motor de Deforestación: Compara la clasificación de cobertura terrestre entre "
            "dos períodos (T1 vs T2). Se identifican transiciones de clases forestales (bosque denso, "
            "manglar, matorral) hacia clases no forestales (pastizal, suelo, urbano).",
            "Motor de Expansión Urbana: Detecta transiciones hacia la clase 'urbano' desde "
            "cualquier otra cobertura. Incluye alertas automáticas para posibles construcciones "
            "sin permiso.",
            "Motor de Vegetación: Clasificación de 11 clases de cobertura terrestre basada "
            "en Dynamic World V1. Se calcula el porcentaje de cada clase sobre el AOI.",
            "Hansen Global Forest Change (v1.12): Datos históricos de pérdida forestal del "
            "dataset UMD/Hansen a resolución 30m. Cubre el periodo 2001-2024 y utiliza bandas de "
            "treecover2000, loss y lossyear para vectorizar zonas de pérdida con filtro de bosque ≥30%.",
            "Alertas GLAD-S2 / RADD: GLAD-S2 utiliza Sentinel-2 (óptico) para detección de "
            "deforestación confirmada. RADD utiliza Sentinel-1 (SAR) que penetra nubes — crítico para "
            "regiones tropicales como Chiapas, Oaxaca y Veracruz.",
            "Drivers de Deforestación (WRI): Clasificación de causas de deforestación a 1km "
            "desarrollada por WRI / Google DeepMind. Incluye: agricultura permanente, commodities, "
            "cultivo rotacional, tala, incendios, asentamientos e infraestructura.",
            "Incendios (MODIS MCD64A1): Áreas quemadas detectadas por MODIS a resolución 500m. "
            "Se correlaciona con polígonos de deforestación para identificar causalidad.",
            "Puntos de Calor FIRMS: Detecciones de fuego activo en tiempo casi real del sistema "
            "NASA FIRMS con sensores VIIRS (375 m, satélites S-NPP y NOAA-20) y MODIS (1 km, Terra y Aqua). "
            "Se reporta FRP (Fire Radiative Power) en megawatts como indicador de intensidad. "
            "Datos locales indexados desde 2018 con respaldo de API en línea.",
            "Contexto Legal (WDPA): Intersección del AOI con Áreas Naturales Protegidas (ANPs) "
            "de la CONANP/WDPA. La deforestación dentro de una ANP constituye infracción potencial "
            "bajo el Artículo 47 de la LGEEPA.",
            "Detección de Anomalías: Se calcula el z-score de cada indicador a lo largo de la "
            "serie temporal. Valores con z > 2σ se clasifican como anomalías de severidad 'media' "
            "(2-3σ) o 'alta' (>3σ).",
            "Nivel de Riesgo: Score compuesto (0-100) que pondera: deforestación acumulada, "
            "expansión urbana, cambio en cobertura forestal, número de anomalías y tasa anual.",
        ]
        for t in texts:
            doc.add_paragraph(t, style="APEXBody")

    # ── PUBLIC API ──
    def generate(self, summary: dict, job_id: str,
                 aoi_geojson: dict = None,
                 thumbnails: dict = None,
                 overview_thumb: dict = None) -> io.BytesIO:
        folio = f"PROFEPA-APEX-{job_id[:8].upper()}"
        cumulative = summary.get("cumulative", {})
        anomalies = summary.get("anomalies", [])
        timeline = summary.get("timeline", {})
        period = cumulative.get("period", "N/D")
        interp = _generate_interpretation(timeline, cumulative, anomalies)
        thumbnails = thumbnails or {}

        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        self._add_styles(doc)
        self._add_cover(doc, folio, period, interp)

        # ── NEW: Location overview map ──
        if aoi_geojson:
            try:
                coords = APEXPDFReportGenerator._compute_coords(aoi_geojson)
                area_ha = _to_float(cumulative.get("total_area_ha",
                            cumulative.get("area_ha", 0)))
                doc.add_paragraph("Ubicación Geográfica", style="APEXHeading")
                loc_buf = _build_location_map(aoi_geojson, coords, area_ha,
                                              overview_thumb)
                doc.add_picture(loc_buf, width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                doc.add_page_break()
            except Exception:
                pass

        self._add_summary(doc, cumulative, len(anomalies), interp)
        self._add_legal_context(doc, summary)
        self._add_hansen_section(doc, summary)
        self._add_drivers_section(doc, summary)
        self._add_fire_section(doc, summary)
        self._add_deforestation_alerts_section(doc, summary)
        self._add_sar_section(doc, summary)
        self._add_crossval_section(doc, summary)
        self._add_firms_section(doc, summary)
        self._add_alerts(doc, anomalies)
        if timeline:
            self._add_annual_data(doc, timeline)
            self._add_charts(doc, timeline)
            self._add_year_pages(doc, timeline, interp,
                                 aoi_geojson=aoi_geojson,
                                 thumbnails=thumbnails)
        self._add_conclusions(doc, interp)
        self._add_methodology(doc)

        # Footer
        doc.add_paragraph("")
        p = doc.add_paragraph(style="APEXSmall")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.add_run("— Fin del reporte — Generado automáticamente por el sistema APEX —")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf
